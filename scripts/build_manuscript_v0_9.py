#!/usr/bin/env python3
"""Build the v0.9 submission-ready manuscript DOCX from the v0.8 draft."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "paper/submission/ReduLink_Deduplex_QUIC_full_draft_v0_8_peer_review_strengthened.docx"
OUT = ROOT / "paper/submission/ReduLink_full_draft_v0_9_submission_ready.docx"
SUMMARY = ROOT / "results/target_class_warm_update_summary.csv"
FIGURE = ROOT / "figures/target_class/redulink_vs_baseline_warm_update.png"


def norm(text: str) -> str:
    return " ".join(text.split())


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def find_paragraph(doc: Document, startswith: str):
    for paragraph in doc.paragraphs:
        if norm(paragraph.text).startswith(startswith):
            return paragraph
    raise ValueError(f"paragraph not found: {startswith}")


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    new = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new._p)
    if style:
        new.style = style
    set_text(new, text)
    return new


def remove_block(block) -> None:
    element = block._element
    element.getparent().remove(element)


def insert_table_after(paragraph, rows: list[list[str]], style: str = "Table Grid"):
    doc = paragraph.part.document
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = style
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
    paragraph._p.addnext(table._tbl)
    return table


def compact_table(table, font_size: int = 7) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def replace_table_after_caption(doc: Document, caption_prefix: str, rows: list[list[str]]) -> None:
    caption = find_paragraph(doc, caption_prefix)
    for table in list(doc.tables):
        prev = table._tbl.getprevious()
        if prev is not None and prev == caption._p:
            remove_block(table)
            break
    table = insert_table_after(caption, rows)
    compact_table(table)


def replace_figure_after_caption(doc: Document, caption_prefix: str, figure_path: Path) -> None:
    caption = find_paragraph(doc, caption_prefix)
    for paragraph in list(doc.paragraphs):
        prev = paragraph._p.getprevious()
        if prev is not None and prev == caption._p and paragraph._p.xpath(".//w:drawing"):
            remove_block(paragraph)
            break
    fig = insert_paragraph_after(caption, "")
    fig.add_run().add_picture(str(figure_path), width=Inches(6.35))


def summary_rows() -> list[dict[str, str]]:
    with SUMMARY.open(newline="") as fh:
        return list(csv.DictReader(fh))


def f3(value: str) -> str:
    return f"{float(value):.3f}x"


def target_table_rows() -> list[list[str]]:
    rows = [[
        "Target",
        "Best compression",
        "Fixed-block",
        "ReduLink fixed",
        "ReduLink CDC",
        "Interpretation",
    ]]
    for row in summary_rows():
        rows.append([
            row["target"],
            f"{row['best_compression_method']} {f3(row['best_compression_multiplier'])}",
            f3(row["fixed_block_multiplier"]),
            f3(row["redulink_fixed_multiplier"]),
            f3(row["redulink_cdc_multiplier"]),
            row["interpretation"],
        ])
    return rows


def main() -> None:
    doc = Document(SRC)

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in header.paragraphs:
                if "ReduLink / Deduplex-QUIC" in paragraph.text:
                    set_text(paragraph, paragraph.text.replace("ReduLink / Deduplex-QUIC", "ReduLink"))

    set_text(
        doc.paragraphs[0],
        "ReduLink: Authenticated Redundancy-Suppressed Transmission for Effective Reconstructed Throughput over Encrypted WANs",
    )
    set_text(
        doc.paragraphs[1],
        "Michél Nguyen Independent Researcher | minh.systems Draft version 0.9 submission-ready evidence, artifact, and protocol-scope revision, June 2026",
    )
    set_text(
        find_paragraph(doc, "Abstract."),
        "Abstract. Ethernet and wide-area transport protocols are typically evaluated by physical or nominal line rate, yet many real traffic classes contain repeated objects, pages, layers, templates, or versioned artifacts. ReduLink is a representation-layer model and candidate Deduplex-QUIC integration profile for authenticated reference substitution between cooperative encrypted endpoints. Rather than increasing physical line rate, ReduLink can increase effective reconstructed payload throughput when a receiver validates compact references against an epoch-scoped dictionary. This draft specifies the ReduLink abstraction, candidate QUIC integration semantics, sender and receiver state machines, security properties, privacy modes, and conservative accounting rules. The accompanying artifact validates byte-exact reconstruction, fail-closed reference handling, public-fixture benchmarks, deterministic target-class fixtures, a fixed-block reuse approximation baseline, local wall-clock/RSS cost columns, and a TCP endpoint-reconstruction prototype. The artifact is not a QUIC implementation and does not validate production cryptography, replay windows, congestion fairness, 0-RTT, migration, or cross-tenant privacy enforcement.",
    )

    replacements = {
        "A concrete Deduplex-QUIC frame design": "A candidate Deduplex-QUIC integration profile with FULL, REF, MISS, and DICT_ACK behavior, plus negotiated transport parameters for dictionary scope, expansion bounds, 0-RTT policy, and reference policy.",
        "A conservative throughput model for 1 Gbit/s links": "A conservative throughput model and CSV-backed evidence matrix showing both useful and weak target classes.",
        "A reproducible artifact package with fixed and content-defined chunking": "A reproducible artifact package with fixed and content-defined chunking, deterministic target-class fixtures, public-corpora fetch scripts, fixed-block reuse baseline, local wall-clock/RSS cost columns, socket prototype, automated tests, CI, CSV outputs, and generated figures.",
        "5. Deduplex-QUIC instantiation": "5. Candidate Deduplex-QUIC profile",
        "The most deployable public-WAN instantiation is Deduplex-QUIC.": "One candidate public-WAN profile is Deduplex-QUIC.",
        "redulink_transport_parameter = { version: 1, max_dictionary_bytes: 64 MiB, max_chunk_bytes: 64 KiB, target_chunk_bytes: 8 KiB, max_reference_expansion: 256, privacy_mode: per_connection | per_origin, fallback_required: true }": "redulink_transport_parameter = { version: 1, max_dictionary_bytes: 64 MiB, max_chunk_bytes: 64 KiB, target_chunk_bytes: 8 KiB, max_reference_expansion: 64, max_pending_reconstructed_bytes: peer flow-control bounded, max_ref_wait: local policy, privacy_mode: per_connection | per_origin, fallback_required: true, zero_rtt_references: false }",
        "9.5 Public artifacts and baseline comparisons": "9.5 Target-class evidence and baseline comparisons",
        "9.6 Automated tests, reproducible commands, and generated plots": "9.6 Public fixtures, tests, reproducible commands, and generated plots",
        "Table 6. v0.8 evidence excerpt from generated CSV outputs: evidence level, fixed-block reuse approximation, corrected public-corpora rows, and local cost columns.": "Table 6. v0.9 target-class evidence matrix from generated CSV outputs. The table includes weak and negative cases rather than only positive examples.",
        "Figure 6. Regenerated v0.8 synthetic-suite figure: effective reconstructed throughput multiplier by workload and method, including gzip, zstd, fixed-block reuse, ReduLink fixed, and ReduLink CDC.": "Figure 6. v0.9 warm/update target-class evidence: fixed-block reuse versus ReduLink fixed and ReduLink CDC. Values near or below 1x are expected for weak and negative controls.",
        "The v0.8 reproducibility package contains": "The v0.9 reproducibility package contains the runnable model, compatibility wrapper, tests, GitHub Actions workflow, synthetic CSV, public-corpora manifest and result CSV, target-class manifest and result CSV, target-class summary CSV, fixed-block reuse approximation, benchmark scripts with local wall-clock/RSS cost columns, plot-generation scripts, regenerated figures, minimal socket prototype, threat model, and RFC-style protocol appendix.",
        "Core files include src/redulink_proto_v0_5.py": "Core files include src/redulink_model.py, the compatibility wrapper src/redulink_proto_v0_5.py",
        '[14] M. Nguyen, "ReduLink / Deduplex-QUIC repository and reproducibility package," GitHub repository, 2026. Available: https://github.com/pinkysworld/redulink-deduplex-quic': '[14] M. Nguyen, "ReduLink repository and reproducibility package," GitHub repository, version 0.9, 2026. Available: https://github.com/pinkysworld/redulink-deduplex-quic',
    }

    for paragraph in doc.paragraphs:
        text = norm(paragraph.text)
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                set_text(paragraph, replacement)
                break

    set_text(
        find_paragraph(doc, "Loss recovery must not assume"),
        "Loss recovery must not assume that sender and receiver dictionaries are perfectly synchronized. QUIC packet retransmission handles lost frame bytes, but once a REF arrives and fails dictionary resolution, resending the REF does not repair the semantic miss. The receiver blocks delivery and emits MISS; the sender repairs with a new REDULINK_FULL for the same stream id, reconstructed offset, original length, and chunk id.",
    )
    set_text(
        find_paragraph(doc, "The evaluation is separated into evidence levels"),
        "The central evidence for v0.9 is the target-class suite in results/target_class_suite.csv. These deterministic generated fixtures are not production traces; they are controlled positive, weak, and negative cases that expose when warm dictionary references work. The table is deliberately honest: ReduLink loses on generated software-update and log-like cases, is modest on git-like data, is strong on VM/backup-like aligned pages, and remains near 1x on random and independent compressed negative controls.",
    )
    set_text(
        find_paragraph(doc, "The baseline table compares raw bytes"),
        "The fixed-block reuse baseline is intentionally strong. It approximates delta-transfer behavior over exact block matches and is not wire-compatible rsync, but it is a serious comparator for update-like artifacts. It often beats ReduLink when object boundaries remain aligned, because exact block reuse avoids CDC cost and ReduLink pays FULL/REF framing overhead. ReduLink's narrower claim is not dominance over delta transfer; it is authenticated endpoint-controlled reference substitution with transport-aware dictionary, miss, privacy, and accounting semantics.",
    )
    set_text(
        find_paragraph(doc, "The repository now includes tests for byte-exact"),
        "The repository now includes tests for byte-exact FULL/REF reconstruction, random-data negative controls, warm-dictionary gains, safe REF-miss failure, public-manifest validation that skips cleanly until optional corpora are fetched, socket prototype byte-count accounting, and both fixed and content-defined chunkers. GitHub Actions runs unit tests, public-fixture fetch, target-class generation, consistency checks, benchmark smoke tests, figure generation, evidence-table generation, and the socket demo.",
    )
    set_text(
        find_paragraph(doc, "Benchmark reproduction is command-based"),
        "Benchmark reproduction is command-based: python3 benchmarks/fetch_public_corpora.py fetches pinned public pairs; bash benchmarks/run_synthetic_suite.sh emits results/synthetic_suite.csv; bash benchmarks/run_public_artifacts.sh --manifest benchmarks/public_artifacts_manifest.csv emits results/public_artifact_suite.csv; bash benchmarks/run_target_class_suite.sh emits results/target_class_suite.csv; python3 benchmarks/check_generated_artifacts.py checks generated target manifests and result labels; python3 scripts/summarize_benchmark_evidence.py regenerates paper/evidence_tables.md; and python3 scripts/plot_warm_update_summary.py regenerates the paper-facing Figure 6.",
    )
    set_text(
        find_paragraph(doc, "The scope of claim is deliberately narrow"),
        "ReduLink is not a faster link, a universal accelerator, a compression replacement, a delta-transfer replacement, or a completed QUIC implementation. The contribution is a scoped protocol model for authenticated reference substitution in cooperative endpoints, with explicit reconstruction invariants, miss repair, dictionary privacy boundaries, expansion limits, and reproducible evidence showing both useful and weak target classes.",
    )
    set_text(
        find_paragraph(doc, "The current evaluation separates analytic modeling"),
        "The current evaluation separates analytic modeling, deterministic target-class fixtures, small frozen public-corpora fixtures, and a TCP endpoint prototype. It still requires larger public corpora and real QUIC implementation experiments before production-scale claims.",
    )

    replace_table_after_caption(doc, "Table 6.", target_table_rows())
    replace_figure_after_caption(doc, "Figure 6.", FIGURE)

    # Add explicit public-corpus limitation paragraph immediately after Table 6.
    caption = find_paragraph(doc, "Table 6.")
    public_note = insert_paragraph_after(
        caption,
        "The current public fixture remains intentionally small: pinned text/version pairs from CPython, nginx, Redis, Linux documentation, and IETF QUIC RFCs. It is useful as a checksum-verifiable smoke fixture, not as production trace validation. Larger OCI layers, git packs, package repository metadata, VM/backup snapshots, and structured log archives remain required for stronger venue claims.",
    )
    # Place note after the table, not between caption and table.
    for table in doc.tables:
        prev = table._tbl.getprevious()
        if prev is not None and prev == caption._p:
            table._tbl.addnext(public_note._p)
            break

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
