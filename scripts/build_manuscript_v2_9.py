from __future__ import annotations
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import csv, json

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT/'paper/submission/ReduLink_journal_ready_v2_9.docx'


def read_csv(name):
    with (ROOT/'results'/name).open(newline='') as fh:
        return list(csv.DictReader(fh))

def read_json(name):
    with (ROOT/'results'/name).open() as fh:
        return json.load(fh)

def fmtx(v):
    try: return f"{float(v):.2f}x"
    except Exception: return str(v)

def fmtn(v):
    try: return f"{int(float(v)):,}"
    except Exception: return str(v)

def f1(v):
    try: return f"{float(v):.1f}"
    except Exception: return str(v)

def f3(v):
    try: return f"{float(v):.3f}"
    except Exception: return str(v)

def add_cell_text(cell, text, size=8, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.font.size = Pt(size); r.bold = bold
    return cell

def add_table(doc, headers, rows, widths=None, font_size=8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        add_cell_text(t.rows[0].cells[i], h, size=font_size, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            add_cell_text(cells[i], val, size=font_size)
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            tcPr = cell._tc.get_or_add_tcPr(); tcMar = OxmlElement('w:tcMar')
            for m in ['top','left','bottom','right']:
                mar = OxmlElement(f'w:{m}'); mar.set(qn('w:w'), '40'); mar.set(qn('w:type'), 'dxa'); tcMar.append(mar)
            tcPr.append(tcMar)
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph('')
    return t

def para(doc, text, size=10, after=4, before=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text); r.font.size = Pt(size); return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(); p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(2)
    p.add_run('- ' + text).font.size = Pt(size); return p

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = None; r.font.name = 'Arial'
    return p

def tcap(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text); r.font.size = Pt(9); r.bold = True; return p

def figcap(doc, text):
    return para(doc, text, size=9, after=6)

def add_page_numbers(doc):
    sec = doc.sections[0]; fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2); run.font.size = Pt(9)

def picture(doc, relpath, width, cap):
    try:
        doc.add_picture(str(ROOT/relpath), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        figcap(doc, cap)
    except Exception:
        pass

# ---------------- load results ----------------
journal = read_csv('journal_workload_suite.csv')
def mult_for(artifact, method):
    wanted_chunker=None; real=method
    if method=='ReduLink fixed': real='redulink'; wanted_chunker='fixed'
    elif method=='ReduLink CDC': real='redulink'; wanted_chunker='cdc'
    for r in journal:
        if r.get('artifact')!=artifact or r.get('method')!=real: continue
        if wanted_chunker and r.get('chunker')!=wanted_chunker: continue
        if r.get('mode') not in ('warm-update-like','single-object'): continue
        return float(r['effective_multiplier'])
    return None
def bestcomp(artifact):
    vals=[mult_for(artifact,m) for m in ['gzip-6','zstd-3']]
    vals=[v for v in vals if v]; return max(vals) if vals else 0

ext_public = read_csv('external_public_suite.csv')
rsync_public = {r['label']: r for r in read_csv('rsync_baseline_external_public.csv')}
ext_pos = read_csv('external_positive_suite.csv')
ext_obj = read_csv('external_object_workload_suite.csv')
quic_cases = read_csv('aioquic_workload_cases.csv')
quic_flow = read_csv('quic_flow_comparison.csv')
repeat = read_csv('repeated_quic_trials_summary.csv')
block = read_csv('journal_block_size_sensitivity.csv')
comp = read_csv('component_performance.csv')
scaling = read_csv('aioquic_scaling_experiment.csv')
bottleneck = read_csv('quic_bottleneck_emulation.csv')
compflow = read_json('quic_competing_flows.json')
wirefair = read_json('wire_fairness_accounting.json')
semrepair = read_json('semantic_repair_demo.json')
udprepair = read_json('udp_repair_experiment.json')
authudp = read_json('authenticated_udp_experiment.json')

# ---------------- document ----------------
doc = Document()
sec = doc.sections[0]
for m in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec, m, Inches(0.6))
styles = doc.styles
styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10)
for s in ['Title','Heading 1','Heading 2','Heading 3']:
    if s in styles: styles[s].font.name='Arial'; styles[s].font.color.rgb=None
add_page_numbers(doc)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('ReduLink: Authenticated Reference Substitution for Redundancy-Suppressed Transfers over Encrypted QUIC Streams')
r.bold=True; r.font.size=Pt(15)
a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.add_run('Michél Nguyen\nUniversity of the People | minh.systems | ORCID: 0000-0001-6834-4422').font.size=Pt(10)

heading(doc,'Abstract',1)
para(doc,"Encrypted transports such as QUIC deliberately hide payload bytes from the network, which disables the transparent in-network redundancy elimination that wide-area optimizers have historically relied on. ReduLink is an endpoint-controlled representation layer that recovers redundancy savings without breaking encryption: it replaces repeated payload chunks with compact references that are individually authenticated against epoch, scope, stream identifier, reconstructed offset, length, nonce, and chunk identity. The contribution is not a new chunk-matching algorithm and not a claim of faster physical links; it is authenticated, scoped, QUIC-compatible reference substitution with explicit reconstruction, privacy scoping, and fail-closed semantics.")
para(doc,"The artifact maps compact binary ReduLink messages into TLS-protected aioquic QUIC streams, implements semantic MISS/FULL repair, derives keys with an exporter-style schedule, and tests replay and tamper rejection. We evaluate against fixed-block reuse, gzip/zstd, and real rsync over deterministic journal fixtures, hash-pinned public source releases, object-aligned public-release transfers, and a layer-like positive case, and we report native QUIC stream experiments under datagram loss, scaling, repeated trials, block-size sensitivity, component costs, competing-flow and bottleneck-emulation fairness, and conservative wire-byte accounting.")
para(doc,"Quantitatively, on object-aligned public release transfers ReduLink reaches 1.94x to 9.70x effective stream-payload reduction (and 21.37x on a layer-like positive case), within a few percent of an idealized fixed-object-reuse baseline, while the authenticated variant stays within a few additional percent; on ordinary source-tree updates rsync wins by up to 73x, which we report rather than hide. Native aioquic experiments confirm byte-exact reconstruction with deterministic stream multipliers under zero and periodic datagram loss, and the accounting experiments confirm that congestion and fairness are charged on encoded bytes, not reconstructed bytes. The result is a clear deployment boundary: simple reuse or rsync wins for file-tree synchronization, while ReduLink is preferable for object-aligned or layer-like warm-state transfers where an authenticated, stream-compatible reference is the desired abstraction.")
kp=doc.add_paragraph(); kp.paragraph_format.space_before=Pt(2)
kr=kp.add_run('Keywords: '); kr.bold=True; kr.font.size=Pt(9)
kp.add_run('QUIC; HTTP/3; redundancy elimination; data deduplication; content-defined chunking; shared dictionaries; authenticated references; encrypted transport; reproducible evaluation.').font.size=Pt(9)

heading(doc,'1. Introduction',1)
para(doc,"Encrypted transports restrict transparent in-network redundancy elimination. QUIC combines UDP transport, TLS security, stream multiplexing, loss recovery, and congestion control, which makes middlebox rewriting of payloads the wrong abstraction for most public or end-to-end encrypted deployments [6-10]. A WAN optimizer can only suppress redundancy that it can see, and a correctly deployed QUIC connection ensures that it sees only ciphertext. The redundancy, however, has not disappeared: registries, content-delivery networks, software-update services, and backup systems still send byte-stable objects to receivers that already hold closely related bytes from a previous transfer.")
para(doc,"The practical question is therefore not whether repeated bytes can be replaced by references. Fixed-block reuse, rsync, and low-bandwidth file systems already demonstrate that repeated bytes can be avoided [1-5], and HTTP delta and shared-dictionary mechanisms show that receiver-held bytes can be reused as a dictionary on the web [19-22]. The question this paper addresses is whether a reference-substitution mechanism can be made explicit, individually authenticated, privacy-scoped, and compatible with encrypted endpoint streams, without claiming a universal accelerator and without asking the application to become a file-tree synchronization protocol. Figure 1 shows the resulting architecture: cooperating endpoints that already control the plaintext decide, per chunk, whether to send bytes or an authenticated reference, and the receiver reconstructs or fails closed.")
para(doc,"We use one accounting vocabulary throughout. The effective multiplier is reconstructed or input bytes divided by encoded bytes at a stated accounting layer. Stream-payload multipliers exclude UDP/IP/link overhead; UDP-estimated multipliers add a local IPv4/UDP estimate; and congestion fairness, where claimed, applies to encoded bytes, not reconstructed bytes. This separation is deliberate: ReduLink changes how many application bytes a given number of wire bytes can reconstruct, not the physical line rate.")
para(doc,"This paper investigates three questions. RQ1: can reference substitution over encrypted QUIC streams be made explicit, authenticated, and privacy-scoped without custom transport-layer changes? RQ2: under which workload shapes does authenticated reference substitution save bytes, and where do simpler fixed-block reuse, compression, or rsync win? RQ3: what is the component-level cost of the authentication and accounting machinery, and how stable and fair are the native QUIC results across repeated runs, scales, and emulated bottlenecks? Section 7 answers RQ2, Sections 8 and 10 answer RQ3, and Sections 4 and 5 answer RQ1.")

heading(doc,'2. Contributions and Claim Boundary',1)
for item in [
 "An authenticated ReduLink reference format and validation model for endpoint-controlled encrypted streams, with per-frame binding of epoch, scope, stream id, reconstructed offset, length, nonce, and chunk identity.",
 "A compact binary mapping carried inside native aioquic QUIC streams, with byte-exact semantic MISS/FULL repair, repeated trials, and a payload-scaling experiment.",
 "An exporter-style artifact key schedule and security tests for wrong secret, scope, epoch, stream id, offset, length, replay, and tampering, including authenticated-UDP tamper and replay probes.",
 "A comparison against fixed-block reuse, gzip/zstd, and real rsync, including cases where these baselines win decisively, plus an explicit contrast with HTTP delta and shared-dictionary transport.",
 "External public source-release negative evidence, object-aligned public source-release workloads, and a public Redis-derived layer-like positive case, all hash-pinned and reproducible.",
 "Block-size sensitivity, component-cost measurements, competing-flow and bottleneck-emulation fairness analysis, and conservative accounting-layer separation.",
]:
    bullet(doc,item)
para(doc,"The central claim is conditional. ReduLink helps when byte-stable chunks recur across warm endpoint state and authenticated reference substitution is preferable to a file-oriented delta protocol or local compression. It is not a replacement for rsync or compression, and the artifact is a native QUIC stream mapping rather than a custom QUIC extension-frame implementation. A practical example is a registry, CDN, backup, or update service that repeatedly sends object records to the same endpoint or tenant while preserving end-to-end QUIC encryption. In that setting the server can reference receiver-known objects or chunks without exposing plaintext to a middlebox; the price is authentication and metadata overhead, and the benefit is explicit context binding, safe miss repair, and a stream-compatible representation.")

heading(doc,'3. Background and Related Work',1)
heading(doc,'3.1 Network redundancy elimination',2)
para(doc,"Spring and Wetherall introduced protocol-independent redundancy elimination, suppressing repeated byte ranges on a link by caching recently seen content at both ends [1]. Subsequent work generalized this to network-wide and coordinated redundancy elimination (SmartRE) and to enterprise and end-system services (EndRE), and measured how much redundancy real traffic contains [2-4]. These systems are powerful but assume a vantage point that can observe or terminate plaintext. ReduLink targets the opposite regime, where the transport is end-to-end encrypted and only the endpoints can act, so the redundancy decision must move into the endpoint stack.")
heading(doc,'3.2 File-delta and low-bandwidth file systems',2)
para(doc,"rsync and the Low-Bandwidth File System (LBFS) show the value of transferring only the differences between file objects that both sides can name and compare [5]. They are the right tool when the workload is genuinely a file tree and both endpoints can run a delta protocol. As our negative results in Section 7.1 confirm, ordinary source-tree updates are better served by rsync than by stream-level reference substitution. ReduLink is not aimed at this case; it is aimed at object or stream delivery where a file-tree delta negotiation is not the serving abstraction.")
heading(doc,'3.3 HTTP delta and shared-dictionary transport',2)
para(doc,"The web has its own lineage of receiver-side reuse over an encrypted channel, and it is the closest existing work. RFC 3229 defines delta encoding in HTTP, sending only the difference of a response relative to a prior instance the client holds, typically using the VCDIFF format [19,20]. Shared Dictionary Compression over HTTP (SDCH) generalized this to a shared dictionary referenced across responses [21], and the recent Compression Dictionary Transport standard (RFC 9842) lets a server designate a previously fetched response as an external Brotli or Zstandard dictionary for future responses [22]. Because these mechanisms operate on HTTP responses, they already run over HTTP/3 on QUIC: receiver-held bytes are reused as a dictionary across an encrypted transport, which is precisely ReduLink's setting.")
para(doc,"ReduLink differs in three ways that matter for the deployments it targets. First, it authenticates each reference individually and binds it to epoch, scope, stream, reconstructed offset, length, and a nonce, so a reference cannot be replayed across context or silently expanded; compression-dictionary schemes treat the dictionary as a codec input and rely on the surrounding TLS channel for integrity. Second, ReduLink defines an explicit fail-closed semantic-repair path (MISS then FULL) rather than degrading to a codec fallback. Third, it is a general stream-representation abstraction rather than a codec scoped to HTTP responses with URL-pattern dictionary matching. We therefore treat compression dictionary transport as the strongest deployed point of comparison and position ReduLink as the authenticated, fail-closed, transport-general variant of the same receiver-side reuse idea.")
heading(doc,'3.4 Content-defined chunking and deduplication privacy',2)
para(doc,"Content-defined chunking determines where chunk boundaries fall and strongly affects deduplication ratio and throughput; recent surveys and fast designs analyze the trade-offs in depth [11,12]. Chunking and deduplication also create content-existence side channels: deduplication can leak whether a receiver already holds particular content, and chunk-boundary parameters can themselves be attacked [13,16,17]. These results directly motivate ReduLink's privacy scoping (Section 4) and its exclusion of global cross-user dictionaries. Modern Ethernet line rates continue to increase [14], which is why the paper avoids physical-rate claims and reports only representation-layer savings.")
heading(doc,'3.5 Deployment model',2)
para(doc,"The deployment model has three roles: a sender with access to the new byte stream, a receiver with a scoped warm dictionary, and a policy domain that defines whether dictionary state is per connection, per origin, or per tenant. Public Internet mode defaults to per-connection dictionaries. Public artifact mode may use signed or manifest-controlled per-origin dictionaries. Enterprise mode may use per-tenant dictionaries with audit and quota controls. Global private cross-user dictionaries remain out of scope. Table 1 summarizes when ReduLink is and is not the right tool relative to rsync and compression.")
tcap(doc,'Table 1. ReduLink deployment modes and dictionary scoping.')
add_table(doc, ['Deployment','Dictionary scope','Why ReduLink rather than rsync/compression'], [
    ['CDN or registry objects','Per origin or per client','Objects arrive as encrypted streams; file-tree delta negotiation may not be the serving abstraction.'],
    ['Backup/page streams','Per endpoint or tenant','Page-like chunks recur across snapshots and can be referenced while preserving authenticated reconstruction.'],
    ['Enterprise VPN/admin domain','Per tenant','Policy can allow shared warm state while keeping middleboxes away from plaintext.'],
    ['Source-tree sync','File tree','Usually better served by rsync; ReduLink is not targeted here.'],
], widths=[1.4,1.4,3.5], font_size=8)

heading(doc,'4. Protocol and Security Model',1)
para(doc,"ReduLink is a representation layer for cooperating, mutually authenticated endpoints. QUIC packetization, packet numbers, ACK processing, TLS, and congestion control remain QUIC responsibilities; ReduLink defines only how application bytes are represented as authenticated FULL, REF, MISS, and DICT_ACK records and how the receiver validates them. The sender chunks outgoing bytes and emits FULL frames for new chunks or REF frames for chunks already believed present at the receiver. A REF is deliverable only if dictionary membership, epoch, scope, stream id, reconstructed offset, length, nonce, chunk identity, and authentication tag all validate; otherwise the receiver fails closed and requests semantic FULL repair. The full sender/receiver state machines, dictionary-admission rules, and negotiation parameters are specified in the artifact's protocol appendix.")
picture(doc,'figures/architecture/redulink_architecture.png',6.7,"Figure 1. ReduLink protocol and deployment architecture. The sender chunks outgoing bytes and emits FULL or authenticated REF records carried as compact binary application data inside an encrypted QUIC stream; the receiver validates each reference against dictionary membership and bound context before byte-exact reconstruction, failing closed and requesting semantic FULL repair on any mismatch.")
heading(doc,'4.1 Frame types and wire model',2)
para(doc,"ReduLink uses four logical frame types. A FULL frame carries chunk bytes and admits them to the receiver dictionary after validation. A REF frame carries an authenticated reference to a chunk already present at the receiver. A MISS frame is a receiver request for FULL fallback when a reference cannot be resolved. A DICT_ACK frame optionally acknowledges that a receiver admitted a chunk, letting a conservative sender reference only acknowledged chunks. Critically, the stream_offset field always denotes the offset in the original reconstructed application byte stream, never an offset in encoded frame bytes. Table 2 lists the bound fields. The artifact uses conservative accounting overheads of 24 bytes for FULL metadata and 32 bytes for REF metadata; these are accounting inputs, not a final QUIC wire encoding, which would recompute overhead from varint lengths and authentication-tag choice.")
tcap(doc,'Table 2. ReduLink frame types and their authenticated fields.')
add_table(doc, ['Frame','Purpose','Bound / carried fields'], [
    ['FULL','Carry and admit new chunk bytes','epoch, stream id, stream offset, chunk length, chunk id, payload, auth tag'],
    ['REF','Reference a chunk already at the receiver','epoch, stream id, stream offset, original length, chunk id, reference nonce, auth tag'],
    ['MISS','Request FULL fallback (fail-closed)','epoch, stream id, stream offset, chunk id'],
    ['DICT_ACK','Acknowledge dictionary admission','epoch, chunk id, dictionary generation'],
], widths=[1.0,2.3,3.3], font_size=8)
heading(doc,'4.2 Validation and fail-closed reconstruction',2)
para(doc,"The receiver applies a fixed validation order before any byte is delivered, summarized as follows:")
for s in [
 "Reject frames whose epoch does not match the active epoch.",
 "For FULL, validate the authentication tag and chunk identifier, admit the payload to the dictionary, and deliver the bytes at the intended stream offset.",
 "For REF, validate the authentication tag, expansion bound, stream offset, original length, nonce, and dictionary presence; deliver reconstructed bytes only after all checks succeed.",
 "Emit MISS when the referenced chunk is unavailable or policy refuses the reference; a REF that cannot become valid in the current epoch triggers an immediate MISS.",
 "Apply stream ordering and flow control to reconstructed bytes, while applying congestion accounting to transmitted wire bytes.",
]:
    bullet(doc,s)
para(doc,"A reference miss is treated as a synchronization failure, not a corruption event: the receiver delivers no reconstructed bytes, emits MISS, and the sender repairs with a FULL for the same stream id, offset, length, and chunk id. A REF whose original_length disagrees with the stored chunk length, or that would expand beyond the negotiated bound, is a hard validation failure. Section 11 reports prototypes that exercise exactly these miss, tamper, and replay paths.")
heading(doc,'4.3 Dictionaries, epochs, and key schedule',2)
para(doc,"Dictionaries are scoped per connection by default and per origin or per tenant only under explicit policy; global cross-user dictionaries are out of scope. Implementations use bounded LRU-style eviction, per-origin or per-tenant quotas, and short-lived epochs, and a receiver may refuse admission under memory pressure or policy. Chunk identifiers are computed with domain separation over the active epoch, dictionary scope, stream context where applicable, a canonical length encoding, and the chunk bytes, so the same bytes in a different context yield a different identifier. The artifact derives inner authentication keys with an HKDF exporter-style schedule keyed by an explicit experiment secret plus ALPN, epoch, scope, connection, and stream context; tests confirm that the derived secret changes when any of these inputs changes. A production profile should instead key this schedule from QUIC TLS exporter bytes.")
heading(doc,'4.4 Threat model',2)
para(doc,"Table 3 maps each security property to the mechanism it requires and to the current artifact status, and Table 4 names the privacy modes and their leakage risks. The artifact implements HMAC binding, chunk-id validation, nonce rejection, length and expansion checks, and fail-closed repair; it does not implement production exporter-derived keys, custom extension-frame parsing, production replay windows, or cross-tenant isolation enforcement, which are stated as future work. Like other deduplication systems, ReduLink can create a content-existence oracle when an adversary can choose payloads or references and observe transfer size, timing, or MISS behavior; per-connection dictionaries remove the cross-user oracle in public mode, and shared dictionaries are permitted only for public artifacts or explicit trust domains.")
tcap(doc,'Table 3. Security properties, required mechanisms, and artifact status.')
add_table(doc, ['Property','Required mechanism','Artifact status'], [
    ['Integrity','FULL/REF authentication, chunk-id, offset and length binding','Modeled by reconstruction and mismatch tests; production crypto pending.'],
    ['Context binding','Exporter-derived keys, epoch/stream/offset/scope, nonce, replay window','HMAC binding and nonce rejection implemented; exporter keys pending.'],
    ['Dictionary safety','Authenticated FULL, manifest commitment, admission/eviction policy','FULL chunk-id checks modeled; manifest policy not implemented.'],
    ['Expansion bound','Per-frame, per-stream, per-epoch reconstructed-byte caps','Length and accounting checks modeled; full flow control pending.'],
    ['Privacy scope','Per-connection default, no global cross-user dictionary','Policy specified; cross-tenant enforcement not implemented.'],
], widths=[1.2,3.0,2.4], font_size=8)
tcap(doc,'Table 4. Privacy modes and leakage risks.')
add_table(doc, ['Mode','Dictionary scope','Leakage risk','Default'], [
    ['Public Internet','Per connection only','Same-connection access-pattern leakage','Yes'],
    ['Public artifacts','Per-origin signed manifest','Artifact-version / popularity inference','Optional'],
    ['Enterprise VPN','Tenant / admin domain','Intra-tenant content-existence leakage','Optional'],
    ['CDN/update channel','Origin-scoped public versions','Version-possession inference','Optional'],
    ['Global cross-user','Any user','Private content-possession leakage','No (out of scope)'],
], widths=[1.3,1.7,2.4,0.9], font_size=8)

heading(doc,'5. Implementation and Metrics',1)
para(doc,"The artifact is implemented in Python for reproducibility and uses aioquic for native QUIC stream experiments [18]. ReduLink messages are compact binary application-stream records inside QUIC streams; this exercises a real QUIC handshake and encryption but does not implement custom QUIC frame types or transport parameters. The JSON encoding is retained only as a readable baseline, and the default path uses the compact binary format, which encodes frame type, stream id, reconstructed offset, length, chunk id, nonce, authentication tag, and payload only where required.")
para(doc,"Three byte layers are reported. Input bytes are the reconstructed application bytes. Stream-payload bytes are the ReduLink or raw application bytes written into QUIC streams. UDP-estimated bytes add a local IPv4/UDP estimate from the loopback proxy path. The paper avoids treating stream-payload multipliers as full wire-rate measurements; they isolate the representation layer so that raw QUIC, ReduLink, gzip, zstd, fixed reuse, and rsync-style baselines can be compared consistently. Table 5 records, for each evidence layer, what it supports and what it does not prove.")
para(doc,"Two reproducibility notes apply to the absolute numbers. Byte multipliers are deterministic functions of the input bytes and chunking parameters and are therefore machine-independent; they are the paper's results. Wall-clock times and local throughputs (for example the client elapsed times in Section 8 and the MiB/s figures in Section 9) are indicative measurements from a single commodity multi-core Linux host and are hardware-dependent; they should be read as order-of-magnitude component costs, not as portable performance guarantees.")
tcap(doc,'Table 5. Evidence hierarchy: what each layer supports and does not prove.')
add_table(doc, ['Evidence','Supports','Does not prove'], [
    ['Model','FULL/REF accounting and reconstruction','Transport behavior'],
    ['Secure model','HMAC binding and replay checks','Live QUIC exporter use'],
    ['aioquic stream','Encrypted QUIC stream mapping','Custom extension frames'],
    ['UDP/IPv4 estimate','Local datagram accounting','Full packet capture'],
    ['Workloads','Workload sensitivity','Universal acceleration'],
], widths=[1.3,2.6,2.6])

heading(doc,'6. Evaluation Methodology',1)
para(doc,"Evaluation uses four classes of evidence. First, deterministic journal fixtures (disk snapshot, OCI layer, package metadata, repository snapshot, structured logs, and a compressed negative control) isolate predicted positive and negative workload shapes. Second, public source-release pairs (Click, Redis, nginx) test ordinary source-tree updates where ReduLink should not be assumed to help. Third, object-aligned public release experiments use the same public bytes but model registry/CDN object delivery rather than tarball synchronization. Fourth, native aioquic stream experiments measure the encrypted stream mapping on positive and negative cases, under loss, at scale, and against competing flows and emulated bottlenecks.")
para(doc,"For each workload the paper reports ReduLink fixed chunking, ReduLink content-defined chunking where relevant, fixed-block reuse, compression, and rsync where the baseline is structurally applicable. The fixed-block baseline is intentionally strong and simple: if it beats ReduLink, the result is reported rather than hidden, because ReduLink is a security and transport-compatibility layer over reuse, not a superior matching algorithm. All external corpora are hash-pinned, and every figure and table is regenerated from committed result files.")

heading(doc,'7. Workload Results and Baseline Interpretation',1)
para(doc,"Table 6 reports the deterministic warm-state fixtures. On most byte-stable workloads, fixed-block reuse matches or beats ReduLink; the distinct ReduLink contribution is not superior matching but authenticated, scoped, stream-compatible reference substitution with explicit repair and privacy rules. The negative rows (package metadata, logs, and the compressed control) are included precisely because they fail: where repetition is local to one object, compression wins, and where there is no reusable structure, both fixed reuse and ReduLink correctly decline to help.")
def bcomp(a): return bestcomp(a)
artifacts=[('scripted-disk-snapshot','disk snap'),('scripted-oci-layer','oci layer'),('scripted-package-metadata','package meta'),('scripted-repository-snapshot','repository snap'),('scripted-structured-logs','logs'),('independent-compressed-negative','compressed neg')]
rows=[]
for art,label in artifacts:
    rl=mult_for(art,'ReduLink fixed') or mult_for(art,'ReduLink CDC') or 0
    fixed=mult_for(art,'fixed-block-reuse') or 0; bc=bcomp(art)
    rows.append([label, fmtx(rl), fmtx(fixed), fmtx(bc), f"{rl-fixed:+.2f}x"])
tcap(doc,'Table 6. Warm-state workload results (effective stream-payload multiplier): ReduLink versus fixed-block reuse and best compression.')
add_table(doc, ['Workload','ReduLink','Fixed reuse','Best compression','RL - fixed'], rows, widths=[1.5,1.1,1.1,1.3,1.1])

heading(doc,'7.1 External Public Source Releases',2)
para(doc,"Table 7 reports three independent, hash-pinned source-release pairs. The result is negative for ReduLink and strongly positive for rsync. This is an important and deliberately reported finding: ordinary related source trees are better served by file-oriented delta transfer than by stream-level reference substitution, because rsync can exploit fine-grained intra-file similarity that whole-object referencing cannot.")
rows=[]
for r in ext_public:
    rs=rsync_public.get(r['label'],{})
    rows.append([r['label'].replace('-to-',' to '), fmtx(r['redulink_multiplier']), fmtx(r['fixed_block_reuse_multiplier']), fmtx(rs.get('rsync_effective_multiplier_control_plus_data',''))])
tcap(doc,'Table 7. External public source-release pairs: ReduLink and fixed reuse versus real rsync (negative case for ReduLink).')
add_table(doc, ['Public pair','ReduLink','Fixed reuse','rsync total'], rows, widths=[2.6,1.0,1.0,1.0])

heading(doc,'7.2 Object-Aligned Public Release Transfers',2)
para(doc,"Table 8 and Figure 2 use the same public release tarballs but change the transfer abstraction: files are treated as object records and chunked independently, modeling registry/CDN object delivery rather than source-tree synchronization. Redis and nginx show positive ReduLink results because many public release objects remain byte-identical across versions; Click is weaker because fewer files are unchanged. ReduLink tracks the idealized fixed-object-reuse baseline within a few percent, and the authenticated variant costs only a few additional percent. This is a stronger external positive signal than a purely synthetic corpus, though it is still a transfer-model experiment rather than a captured production trace.")
rows=[]
for r in ext_obj:
    label=r['label'].split('-object-sequence')[0]
    files=f"{r['unchanged_file_count']}/{r['new_file_count']} unchanged"
    rows.append([label, files, fmtn(r['input_bytes']), fmtx(r['redulink_multiplier']), fmtx(r['secure_multiplier']), fmtx(r['fixed_object_reuse_multiplier']), fmtx(r['gzip_new_object_stream_multiplier'])])
tcap(doc,'Table 8. Object-aligned public release transfers modeling registry/CDN object delivery.')
add_table(doc, ['Public release','File stability','Bytes','ReduLink','Secure RL','Fixed object reuse','gzip'], rows, widths=[1.2,1.2,0.8,0.8,0.8,1.0,0.8], font_size=7)
picture(doc,'figures/external_object_workload/external_object_workload_multipliers.png',6.4,"Figure 2. Object-aligned public release transfers: effective stream-payload multipliers for ReduLink, the authenticated variant, fixed object reuse, and gzip.")

heading(doc,'7.3 Layer-Like External-Positive Case',2)
para(doc,"Table 9 reports a Redis-derived layer-like case built from included public Redis release bytes and aligned changed blocks. Its role is to test the predicted positive case for layer-like objects where stable chunks remain aligned; it is not a production trace. Here ReduLink reaches 21.37x with the authenticated variant at 18.96x, again close to the fixed-reuse reference.")
r=ext_pos[0]
tcap(doc,'Table 9. Layer-like external-positive case derived from public Redis release bytes.')
add_table(doc, ['Case','Input bytes','ReduLink','Secure RL','Fixed reuse'], [[r['label'], fmtn(r['new_bytes']), fmtx(r['redulink_multiplier']), fmtx(r['secure_multiplier']), fmtx(r['fixed_block_reuse_multiplier'])]], widths=[2.3,1.0,1.0,1.0,1.0])

heading(doc,'7.4 Block-Size Sensitivity',2)
para(doc,"Table 10 and Figure 3 show that the best block size is workload-dependent: smaller blocks capture more aligned reuse but pay more per-chunk overhead, while larger blocks lose alignment. The artifact therefore treats 4 KiB as a reproducible default rather than a universal optimum, and a production deployment would tune the chunk size per workload class.")
arts=['scripted-disk-snapshot','scripted-oci-layer','scripted-repository-snapshot']
rows=[]
for art in arts:
    vals=[]
    for size in ['1024','2048','4096','8192','16384']:
        m=''
        for r in block:
            if r['artifact']==art and r['chunker']=='fixed' and r['chunk_size']==size: m=fmtx(r['effective_multiplier']); break
        vals.append(m)
    rows.append([art.replace('scripted-','').replace('-',' ')]+vals)
tcap(doc,'Table 10. Block-size sensitivity of the effective multiplier (fixed chunking).')
add_table(doc, ['Workload','1 KiB','2 KiB','4 KiB','8 KiB','16 KiB'], rows, widths=[1.5,0.8,0.8,0.8,0.8,0.8])
picture(doc,'figures/block_size/block_size_sensitivity.png',5.6,"Figure 3. Block-size sensitivity of the effective multiplier for fixed chunking across three warm-update fixtures.")

heading(doc,'8. Native QUIC Experiments',1)
heading(doc,'8.1 Stream mapping under loss',2)
para(doc,"Table 11 compares raw and ReduLink stream mapping at zero and periodic datagram loss, and Table 12 reports positive and negative workload cases. ReduLink reconstructs byte-exactly in every case, including the lossy proxy path, and the compressed-negative control correctly yields no gain. Figure 4 plots the workload cases. Because ReduLink writes fewer stream-payload bytes, its UDP-estimated multiplier stays above one on positive workloads even after the IPv4/UDP per-datagram overhead is added.")
rows=[]
for r in quic_flow:
    rows.append([r['method'].replace('-quic-stream','').replace('redulink-binary','ReduLink').replace('raw','Raw'), r['loss_every'], fmtx(r['effective_multiplier']), fmtx(r['approx_ipv4_udp_multiplier_seen']), 'OK' if r['reconstruction_ok']=='True' else 'FAIL'])
tcap(doc,'Table 11. Native aioquic stream mapping: raw versus ReduLink under zero and periodic datagram loss.')
add_table(doc, ['Method','Loss every','Stream x','UDP-est x','Recon.'], rows, widths=[1.8,0.9,0.9,0.9,0.8])
rows=[]
for r in quic_cases:
    rows.append([r['label'].replace('independent-compressed-negative','compressed negative').replace('external-positive-redis-layered','Redis layered').replace('demo-positive','demo positive'), fmtn(r['input_bytes']), fmtx(r['stream_payload_multiplier']), fmtx(r['approx_ipv4_udp_multiplier_seen']), r['semantic_misses'], 'OK' if r['reconstruction_ok']=='True' else 'FAIL'])
tcap(doc,'Table 12. Native aioquic stream mapping across positive and negative workload cases.')
add_table(doc, ['Case','Input','Stream x','UDP-est x','Misses','Recon.'], rows, widths=[1.8,0.9,0.9,0.9,0.7,0.7])
picture(doc,'figures/quic_workload_cases.png',6.4,"Figure 4. Native aioquic stream mapping on positive and negative workload cases.")

heading(doc,'8.2 Repeated trials and scaling',2)
para(doc,"Table 13 reports three repeated trials of the native demo workload. The multiplier is deterministic for fixed bytes, while elapsed times vary modestly across runs; reporting both avoids cherry-picking a single favorable QUIC run. Table 14 then scales the payload from 96 to 1024 blocks: the stream multiplier rises from 3.19x to 3.87x as more reusable chunks accumulate, the miss count grows proportionally with the dictionary warm-up, and reconstruction remains byte-exact at every scale.")
sel=[]
for r in repeat:
    if r['metric'] in {'raw_client_ms','redulink_stream_multiplier','redulink_udp_est_multiplier','redulink_client_ms'}:
        sel.append([r['metric'].replace('_',' '), r['n'], r['mean'], r.get('stdev',''), r['min'], r['max']])
tcap(doc,'Table 13. Repeated native QUIC trials (n = 3): run-to-run stability of multipliers and elapsed time.')
add_table(doc, ['Metric','n','Mean','Std. dev.','Min','Max'], sel, widths=[2.0,0.5,1.0,1.0,1.0,1.0], font_size=7)
rows=[]
for r in scaling:
    rows.append([r['payload_blocks'], fmtn(r['input_bytes']), fmtx(r['stream_payload_multiplier']), r['semantic_misses'], f1(r['client_elapsed_ms']), 'OK' if r['reconstruction_ok']=='True' else 'FAIL'])
tcap(doc,'Table 14. Payload scaling of the native aioquic stream mapping.')
add_table(doc, ['Blocks','Input bytes','Stream x','Misses','Client ms','Recon.'], rows, widths=[1.0,1.2,1.0,0.9,1.0,0.8])

heading(doc,'9. Component Costs',1)
para(doc,"Table 15 reports component-level throughput on the reference host. Fixed chunking, HMAC validation, and compact binary encoding are fast; the Python content-defined chunker is slow and is explicitly not production line-rate evidence. A production implementation would require optimized native chunking and tighter integration with the transport stack. As noted in Section 5, these throughputs are hardware-dependent component costs, not portable guarantees.")
components=['fixed_chunking','cdc_chunking','model_fixed_encode_decode','secure_hmac_encode_roundtrip','binary_wire_encode','binary_wire_decode']
rows=[]
for c in components:
    for r in comp:
        if r['component']==c:
            summary=json.loads(r['result_summary']); ok=summary.get('reconstruction_ok', summary.get('roundtrip_reconstruction_ok',''))
            ok='OK' if ok is True else ('FAIL' if ok is False else 'not applicable')
            label=c.replace('model_fixed_encode_decode','model fixed roundtrip').replace('secure_hmac_encode_roundtrip','HMAC roundtrip').replace('binary_wire_encode','wire encode').replace('binary_wire_decode','wire decode').replace('fixed_chunking','fixed chunking').replace('cdc_chunking','CDC chunking')
            rows.append([label, fmtn(r['input_bytes']), f1(r['throughput_mib_s_local']), ok])
tcap(doc,'Table 15. Component-level cost measurements (local Python artifact timing; not production line-rate).')
add_table(doc, ['Component','Input bytes','MiB/s local','Roundtrip'], rows, widths=[2.1,1.2,1.1,1.2])

heading(doc,'10. Fairness and Accounting',1)
para(doc,"The central fairness rule is that ReduLink must not receive congestion credit for reconstructed bytes: congestion control and bottleneck service account encoded wire bytes only. Table 16 reports the wire-byte accounting experiment, in which ReduLink's encoded wire share is "+f3(wirefair['wire_share_redulink'])+" against the raw stream's "+f3(wirefair['wire_share_raw'])+" over "+str(wirefair['rounds'])+" rounds, while still reconstructing the same application bytes; the effective application multiplier is "+fmtx(wirefair['redulink_effective_app_multiplier'])+". This confirms that the savings appear as fewer encoded bytes, which is the only place a fair transport should credit them.")
tcap(doc,'Table 16. Wire-byte fairness accounting: ReduLink is charged on encoded bytes, not reconstructed bytes.')
add_table(doc, ['Metric','Value'], [
    ['Fairness rule','Congestion and bottleneck service count encoded wire bytes'],
    ['Raw encoded wire share', f3(wirefair['wire_share_raw'])],
    ['ReduLink encoded wire share', f3(wirefair['wire_share_redulink'])],
    ['ReduLink uses less wire than raw', 'yes' if wirefair['redulink_uses_less_wire_than_raw'] else 'no'],
    ['Effective application multiplier', fmtx(wirefair['redulink_effective_app_multiplier'])],
    ['Rounds', str(wirefair['rounds'])],
], widths=[2.6,3.2])
para(doc,"Table 17 reports a concurrent competing-flow smoke test in which a raw QUIC flow and a ReduLink flow share a localhost schedule. Jain's fairness index on encoded rates is "+f3(compflow['encoded_rate_jain_index'])+" and on reconstructed rates is "+f3(compflow['reconstructed_rate_jain_index'])+", with all flows reconstructing correctly; this is a smoke test, not a full congestion-control study. Table 18 then emulates a fluid fair-share bottleneck across rate and RTT points using the measured encoded stream bytes. Because ReduLink transmits fewer encoded bytes, it completes sooner and delivers higher reconstructed goodput at the same encoded fair share; for example at 25 Mbps and 10 ms it completes in "+f1(29.722)+" ms versus "+f1(51.318)+" ms for raw, delivering "+f1(26.459379)+" Mbps of reconstructed goodput. The benefit narrows as the path becomes RTT-bound, as the 50 ms rows show, because latency rather than encoded volume then dominates completion time.")
tcap(doc,'Table 17. Concurrent competing-flow smoke test (Jain fairness index).')
add_table(doc, ['Metric','Value'], [
    ['Scenario', f"{int(compflow['rounds'])} concurrent rounds, ~{f1(compflow['rate_hint_mbps'])} Mbps hint, loss every {compflow['loss_every']}"],
    ['Encoded-rate Jain index', f3(compflow['encoded_rate_jain_index'])],
    ['Reconstructed-rate Jain index', f3(compflow['reconstructed_rate_jain_index'])],
    ['All flows reconstructed', 'yes' if compflow['all_reconstructed'] else 'no'],
    ['Scope', 'localhost concurrent smoke test, not a full congestion-control study'],
], widths=[2.4,3.4], font_size=8)
# bottleneck scenarios
scn={}
for r in bottleneck:
    key=(r['rate_mbps'], r['rtt_ms']); scn.setdefault(key,{})[r['method']]=r
rows=[]
for (rate,rtt),d in scn.items():
    raw=d['raw-quic-stream']; rl=d['redulink-binary-quic-stream']
    rows.append([f1(rate), f1(rtt), f1(raw['completion_ms_emulated']), f1(rl['completion_ms_emulated']), f1(raw['reconstructed_goodput_mbps_emulated']), f1(rl['reconstructed_goodput_mbps_emulated'])])
tcap(doc,'Table 18. Fluid fair-share bottleneck emulation over measured encoded stream bytes (raw versus ReduLink).')
add_table(doc, ['Rate Mbps','RTT ms','Raw compl. ms','RL compl. ms','Raw goodput Mbps','RL recon goodput Mbps'], rows, widths=[1.0,0.8,1.1,1.1,1.2,1.4], font_size=7)
para(doc,"The package does not include a full tc/netem or Mininet congestion-control study; the competing-flow and bottleneck results are deliberately conservative emulations that isolate the encoded-byte accounting principle rather than claim production fairness.")

heading(doc,'11. Repair and Authentication Prototypes',1)
para(doc,"Beyond the in-stream experiments, three prototypes exercise the fail-closed and authentication paths directly. Table 19 summarizes them. The semantic-repair demo forces a 25 percent dictionary-miss fraction and confirms that every miss is repaired with a FULL and that reconstruction is byte-exact. The localhost UDP repair experiment drops every seventh data datagram and uses timeout-based retransmission, again reconstructing fully after "+str(udprepair['repair_full_frames'])+" semantic repairs and "+str(udprepair['client_retransmissions'])+" retransmissions. The authenticated UDP experiment adds explicit negative probes: a tampered authentication tag and a replayed nonce are both rejected ("+str(authudp['tamper_probe_rejections'])+" tamper and "+str(authudp['replay_probe_rejections'])+" replay rejection) before normal authenticated repair traffic is accepted, demonstrating fail-closed behavior under active manipulation.")
tcap(doc,'Table 19. Repair and authentication prototypes (localhost).')
add_table(doc, ['Prototype','Input bytes','Misses / repairs','Negative probes rejected','Recon.'], [
    ['Semantic repair demo', fmtn(semrepair['input_bytes']), f"{semrepair['misses']} / {semrepair['repair_full_frames']}", 'n/a (25% missing)', 'OK' if semrepair['reconstruction_ok'] else 'FAIL'],
    ['UDP repair (drop every 7)', fmtn(udprepair['input_bytes']), f"{udprepair['semantic_misses']} / {udprepair['repair_full_frames']}", f"{udprepair['client_retransmissions']} retransmits", 'OK' if udprepair['reconstruction_ok'] else 'FAIL'],
    ['Authenticated UDP', fmtn(authudp['input_bytes']), f"{authudp['semantic_misses']} / {authudp['repair_full_frames']}", f"{authudp['tamper_probe_rejections']} tamper, {authudp['replay_probe_rejections']} replay", 'OK' if authudp['reconstruction_ok'] else 'FAIL'],
], widths=[1.8,1.0,1.2,1.6,0.7], font_size=8)

heading(doc,'12. Why Not Just Compress or Use rsync?',1)
para(doc,"Compression is best when repetition is local to one object; rsync is best when both endpoints share a file tree or file object and can run a delta protocol; and HTTP delta or shared-dictionary transport is best when the application is HTTP and a codec-level dictionary suffices. The ReduLink use case is different: encrypted endpoint streams where dictionary state is already available at the receiver, the transfer is not naturally a file-tree synchronization session, and each reconstruction step must be individually authenticated and fail closed. The evaluation confirms the boundary. On source-release trees, rsync dominates (Table 7). On package metadata and logs, compression or fixed-block reuse dominates (Table 6). On object-aligned public releases and layer-like byte-stable objects (Tables 8 and 9), ReduLink provides authenticated reference substitution at modest overhead over a simple fixed-reference baseline.")
para(doc,"The motivating cases are therefore not arbitrary file synchronization. They are stream-serving cases in which the endpoint already knows that a receiver has dictionary state but must still authenticate every reconstruction step and preserve encrypted transport semantics. A CDN edge, registry, backup client, or enterprise service may choose this representation because it avoids exposing plaintext to a WAN optimizer and avoids changing the application protocol into a file-tree delta session, while keeping a stronger per-reference integrity guarantee than a codec-level shared dictionary provides.")

heading(doc,'13. Limitations and Future Work',1)
para(doc,"The implementation maps ReduLink into QUIC streams rather than adding custom QUIC extension frames or transport parameters, and it uses an exporter-style key schedule model rather than live QUIC TLS exporter bytes. The public object-sequence and layer-like workloads are derived from real release bytes but are transfer-model experiments rather than captured production traces. The fairness evidence (competing-flow Jain indices and fluid bottleneck emulation in Section 10) is conservative and does not replace a full network-emulator congestion-control study, and all transport experiments run over localhost with a deterministic loss proxy rather than over a real path with diverse RTT, bandwidth, and loss. Absolute timings are hardware-dependent.")
para(doc,"Future work should add custom QUIC frame integration and transport-parameter negotiation, live exporter-derived keys, production replay-window and cross-tenant isolation enforcement, tc/netem or Mininet fairness experiments with competing real flows, packet capture, larger and independently captured registry or backup traces, and a direct head-to-head comparison against Compression Dictionary Transport on the same corpora.")

heading(doc,'14. Reproducibility',1)
para(doc,"The package contains source code, public corpora, generated corpora, benchmark scripts, result CSV/JSON files, figures, a citation checker, and tests [15]. The reviewer-facing smoke command is python3 scripts/run_smoke_validation.py and full validation is python3 scripts/run_full_validation.py. aioquic tests skip gracefully when aioquic is unavailable; installing requirements-dev.txt enables complete QUIC stream validation. Every table and figure in this paper is regenerated from the committed result files by scripts/build_manuscript_v2_9.py and scripts/make_journal_figures_v2_8.py.")

heading(doc,'15. Conclusion',1)
para(doc,"ReduLink is best understood as authenticated, scoped reference substitution for encrypted endpoint streams. Its byte savings are conditional and are often reproduced or exceeded by simpler fixed-block reuse, compression, or rsync; the contribution is to make reference substitution explicit, individually authenticated, privacy-scoped, repairable, and compatible with native QUIC stream transport, and to position it precisely against both in-network redundancy elimination and HTTP shared-dictionary transport. The evidence supports ReduLink for warm-state object-aligned or layer-like transfers, while negative source-release results show where it should not be used.")

heading(doc,'Data and Code Availability',1)
para(doc,"All source code, public-corpus manifests, generated corpora, benchmark scripts, result CSV/JSON files, figures, the citation checker, and the test suite are openly available in the ReduLink repository [15] under the MIT License. Every figure and table is regenerated from the committed result files, so the reported numbers can be reproduced from the artifact. Large external public-release corpora are reconstructed from documented public sources rather than redistributed.")

heading(doc,'References',1)
refs=[
"[1] N. T. Spring and D. Wetherall, 'A protocol-independent technique for eliminating redundant network traffic,' ACM SIGCOMM, 2000.",
"[2] A. Anand, V. Sekar, and A. Akella, 'SmartRE: An architecture for coordinated network-wide redundancy elimination,' ACM SIGCOMM, 2009.",
"[3] A. Anand et al., 'Redundancy in network traffic: Findings and implications,' ACM SIGMETRICS, 2009.",
"[4] B. Aggarwal et al., 'EndRE: An end-system redundancy elimination service for enterprises,' USENIX NSDI, 2010.",
"[5] A. Muthitacharoen, B. Chen, and D. Mazieres, 'A low-bandwidth network file system,' ACM SOSP, 2001.",
"[6] J. Iyengar and M. Thomson, 'QUIC: A UDP-based multiplexed and secure transport,' RFC 9000, IETF, 2021.",
"[7] M. Thomson and S. Turner, 'Using TLS to secure QUIC,' RFC 9001, IETF, 2021.",
"[8] J. Iyengar and I. Swett, 'QUIC loss detection and congestion control,' RFC 9002, IETF, 2021.",
"[9] M. Kuehlewind and B. Trammell, 'Applicability of the QUIC transport protocol,' RFC 9308, IETF, 2022.",
"[10] B. Trammell et al., 'Manageability of the QUIC transport protocol,' RFC 9312, IETF, 2022.",
"[11] W. Xia, X. Zou, Y. Zhou, H. Jiang, C. Liu, D. Feng, Y. Hua, Y. Hu, and Y. Zhang, 'The design of fast content-defined chunking for data deduplication based storage systems,' IEEE Transactions on Parallel and Distributed Systems, vol. 31, no. 9, 2020.",
"[12] M. Gregoriadis, L. Balduf, B. Scheuermann, and J. Pouwelse, 'A thorough investigation of content-defined chunking algorithms for data deduplication,' arXiv:2409.06066, 2024.",
"[13] B. Alexeev, C. Percival, and Y. X. Zhang, 'Chunking attacks on file backup services using content-defined chunking,' arXiv:2504.02095, 2025.",
"[14] IEEE 802.3 Ethernet Working Group, 'IEEE 802.3 Ethernet Working Group active projects,' accessed June 2026.",
"[15] M. Nguyen, 'ReduLink artifact and reproducibility package,' GitHub repository, 2026. https://github.com/pinkysworld/redulink-deduplex-quic",
"[16] D. Harnik, B. Pinkas, and A. Shulman-Peleg, 'Side channels in cloud services: Deduplication in cloud storage,' IEEE Security and Privacy, 2010.",
"[17] M. Bellare, S. Keelveedhi, and T. Ristenpart, 'DupLESS: Server-aided encryption for deduplicated storage,' USENIX Security, 2013.",
"[18] aioquic project contributors, 'aioquic: QUIC and HTTP/3 implementation in Python,' software repository, version 1.3.0, 2026. https://github.com/aiortc/aioquic",
"[19] J. Mogul, B. Krishnamurthy, F. Douglis, A. Feldmann, Y. Goland, A. van Hoff, and D. Hellerstein, 'Delta encoding in HTTP,' RFC 3229, IETF, 2002.",
"[20] D. Korn, J. MacDonald, J. Mogul, and K. Vo, 'The VCDIFF generic differencing and compression data format,' RFC 3284, IETF, 2002.",
"[21] J. Butler, W.-H. Lee, B. McQuade, and K. Mixter, 'A proposal for shared dictionary compression over HTTP (SDCH),' IETF Internet-Draft, 2008.",
"[22] P. Meenan and Y. Weiss, 'Compression dictionary transport,' RFC 9842, IETF, 2025.",
]
for ref in refs:
    para(doc, ref, size=8, after=1)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.author='Michél Nguyen'
doc.core_properties.last_modified_by='Michél Nguyen'
doc.core_properties.subject='ORCID: 0000-0001-6834-4422; University of the People'
doc.core_properties.comments='ReduLink journal-ready v2.9 manuscript; ORCID: 0000-0001-6834-4422; University of the People'
doc.save(OUT)
print('saved', OUT)
