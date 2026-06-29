from __future__ import annotations
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import csv, json, zipfile, shutil, re, hashlib

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT/'paper/submission/ReduLink_journal_ready_v2_8.docx'


def read_csv(name):
    with (ROOT/'results'/name).open(newline='') as fh:
        return list(csv.DictReader(fh))

def fmtx(v):
    try:
        return f"{float(v):.2f}x"
    except Exception:
        return str(v)

def fmtn(v):
    try:
        return f"{int(float(v)):,}"
    except Exception:
        return str(v)

def add_cell_text(cell, text, size=8, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.font.size = Pt(size)
    r.bold = bold
    return cell

def add_table(doc, headers, rows, widths=None, font_size=8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i,h in enumerate(headers):
        add_cell_text(t.rows[0].cells[i], h, size=font_size, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            add_cell_text(cells[i], val, size=font_size)
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m in ['top','left','bottom','right']:
                mar = OxmlElement(f'w:{m}')
                mar.set(qn('w:w'), '40')
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            tcPr.append(tcMar)
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph('')
    return t

def para(doc, text, size=10, after=4, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = None
        r.font.name = 'Arial'
    return p

def tcap(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text); r.font.size = Pt(9); r.bold = True
    return p

def figcap(doc, text):
    return para(doc, text, size=9, after=6)

def add_page_numbers(doc):
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _El
    sec = doc.sections[0]
    foot = sec.footer
    fp = foot.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld1 = _El('w:fldChar'); fld1.set(_qn('w:fldCharType'), 'begin')
    instr = _El('w:instrText'); instr.set(_qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = _El('w:fldChar'); fld2.set(_qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.size = Pt(9)

# Load results
journal = read_csv('journal_workload_suite.csv')
# helper method lookup for warm fixture table
def mult_for(artifact, method):
    wanted_chunker = None
    real_method = method
    if method == 'ReduLink fixed':
        real_method = 'redulink'; wanted_chunker = 'fixed'
    elif method == 'ReduLink CDC':
        real_method = 'redulink'; wanted_chunker = 'cdc'
    for r in journal:
        if r.get('artifact') != artifact or r.get('method') != real_method:
            continue
        if wanted_chunker and r.get('chunker') != wanted_chunker:
            continue
        # Prefer warm/update-like rows where available.
        if r.get('mode') not in ('warm-update-like', 'single-object'):
            continue
        return float(r['effective_multiplier'])
    return None

ext_public = read_csv('external_public_suite.csv')
rsync_public = {r['label']: r for r in read_csv('rsync_baseline_external_public.csv')}
ext_pos = read_csv('external_positive_suite.csv')
ext_obj = read_csv('external_object_workload_suite.csv')
quic_cases = read_csv('aioquic_workload_cases.csv')
quic_flow = read_csv('quic_flow_comparison.csv')
repeat = read_csv('repeated_quic_trials_summary.csv')
block = read_csv('journal_block_size_sensitivity.csv')
comp = read_csv('component_performance.csv')

# Document setup
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.55); sec.left_margin = Inches(0.55); sec.right_margin = Inches(0.55)
styles = doc.styles
styles['Normal'].font.name = 'Arial'; styles['Normal'].font.size = Pt(10)
add_page_numbers(doc)
for s in ['Title','Heading 1','Heading 2','Heading 3']:
    if s in styles:
        styles[s].font.name = 'Arial'
        styles[s].font.color.rgb = None

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('ReduLink: Authenticated Reference Substitution for Redundancy-Suppressed Transfers over Encrypted QUIC Streams')
r.bold = True; r.font.size = Pt(15)
a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r=a.add_run('Michél Nguyen\nUniversity of the People | minh.systems | ORCID: 0000-0001-6834-4422')
r.font.size = Pt(10)

heading(doc, 'Abstract', 1)
para(doc, 'ReduLink is an endpoint-controlled representation layer for encrypted WAN transfers. It replaces repeated payload chunks with compact references authenticated against epoch, scope, stream identifier, reconstructed offset, length, nonce, and chunk identity. The contribution is not a new matching algorithm and not a claim of faster physical links. The contribution is authenticated, scoped, QUIC-compatible reference substitution: a way to use redundancy in warm endpoint state while preserving explicit reconstruction, privacy, and fail-closed semantics.')
para(doc, 'The artifact maps compact binary ReduLink messages into TLS-protected aioquic QUIC streams, implements semantic MISS/FULL repair, uses an exporter-style key schedule model, and tests replay and tamper rejection. Evaluation includes fixed-block reuse, gzip/zstd, real rsync, negative public source-release pairs, object-aligned public source-release transfers, a public Redis-derived layer-like case, repeated native QUIC trials, block-size sensitivity, component-cost measurements, and conservative fairness accounting. Results show that fixed reuse or rsync often wins for file-tree synchronization, while ReduLink is useful for object-aligned or layer-like warm-state transfers where authenticated stream-compatible references are the desired abstraction.')
para(doc, 'Quantitatively, on object-aligned public release transfers ReduLink reaches 1.94x to 9.70x effective stream-payload reduction (and up to 21.37x on a layer-like positive case), within a few percent of an idealized fixed-object-reuse baseline, while the authenticated variant stays within a few additional percent; on ordinary source-tree updates rsync wins by up to 73x, which we report rather than hide. Native aioquic stream experiments confirm byte-exact reconstruction with deterministic stream multipliers under zero and periodic datagram loss.')
kp = doc.add_paragraph(); kp.paragraph_format.space_before = Pt(2)
kr = kp.add_run('Keywords: '); kr.bold = True; kr.font.size = Pt(9)
kp.add_run('QUIC; redundancy elimination; data deduplication; content-defined chunking; authenticated references; encrypted transport; reproducible evaluation.').font.size = Pt(9)

heading(doc, '1. Introduction', 1)
para(doc, 'Encrypted transports restrict transparent in-network redundancy elimination. QUIC combines UDP transport, TLS security, stream multiplexing, loss recovery, and congestion control, which makes middlebox rewriting of payloads the wrong abstraction for most public or end-to-end encrypted deployments [6-10].')
para(doc, 'The practical question is therefore not whether repeated bytes can be replaced by references. Fixed-block reuse, rsync, and low-bandwidth file systems already demonstrate that repeated bytes can be avoided [1-5]. The question is whether a reference-substitution mechanism can be made explicit, authenticated, scoped, and compatible with encrypted endpoint streams without claiming a universal accelerator.')
para(doc, 'Effective multiplier means reconstructed or input bytes divided by encoded bytes at the stated accounting layer. Stream-payload multipliers exclude UDP/IP/link overhead. UDP-estimated multipliers add a local IPv4/UDP estimate. Congestion fairness, where claimed, applies to encoded bytes, not reconstructed bytes.')
para(doc, 'This paper investigates three questions. RQ1: can reference substitution over encrypted QUIC streams be made explicit, authenticated, and privacy-scoped without custom transport-layer changes? RQ2: under which workload shapes does authenticated reference substitution save bytes, and where do simpler fixed-block reuse, compression, or rsync win? RQ3: what is the component-level cost of the authentication and accounting machinery, and how stable are the native QUIC results across repeated runs?')

heading(doc, '2. Contributions and Claim Boundary', 1)
for item in [
    'An authenticated ReduLink reference format and validation model for endpoint-controlled encrypted streams.',
    'A compact binary mapping carried inside native aioquic QUIC streams, with byte-exact semantic repair and repeated trials.',
    'An exporter-style artifact key schedule and security tests for wrong secret, scope, epoch, stream id, offset, length, replay, and tampering.',
    'A comparison against fixed-block reuse, gzip/zstd, and real rsync, including cases where these baselines win decisively.',
    'External public source-release negative evidence, object-aligned public source-release workloads, and a public Redis-derived layer-like positive case.',
    'Block-size sensitivity, component-cost measurements, and conservative accounting-layer separation.'
]:
    p = doc.add_paragraph(style=None); p.style = doc.styles['Normal']; p.paragraph_format.left_indent = Inches(0.2); p.paragraph_format.first_line_indent = Inches(-0.2); p.add_run('- ' + item).font.size = Pt(10)
para(doc, 'The central claim is conditional: ReduLink helps when byte-stable chunks recur across warm endpoint state and authenticated reference substitution is preferable to a file-oriented delta protocol or local compression. It is not a replacement for rsync or compression, and the artifact is a native QUIC stream mapping rather than custom QUIC extension frames.')
para(doc, 'A practical example is a registry, CDN, backup, or update service that repeatedly sends object records to the same endpoint or tenant while preserving end-to-end QUIC encryption. In that setting, the server can reference receiver-known objects or chunks without exposing plaintext to a middlebox and without asking the application to become a file-tree synchronization protocol. The price paid by ReduLink is authentication and metadata overhead; the benefit is explicit context binding, safe miss repair, and a stream-compatible representation.')

heading(doc, '3. Related Work and Motivation', 1)
para(doc, 'Spring and Wetherall introduced protocol-independent redundancy elimination, and later work studied enterprise redundancy, coordinated redundancy elimination, and endpoint redundancy elimination [1-4]. LBFS and rsync-family systems show the value of low-bandwidth file synchronization when both sides operate over file objects [5]. Content-defined chunking and chunking security have a substantial literature [11-13]. Deduplication side channels motivate careful dictionary scoping and avoidance of global cross-user private dictionaries [16,17].')
para(doc, 'ReduLink is motivated by a narrower deployment class: encrypted object delivery over QUIC streams, registry or CDN transfers where the receiver has prior same-origin state, backup-like page streams, and administrative domains where per-origin or per-tenant dictionaries are explicit policy. Modern Ethernet line rates continue to increase [14], so the paper avoids physical-rate claims and focuses only on representation-layer savings.')
para(doc, 'The system boundary is deliberately endpoint-side. Middlebox WAN optimizers can exploit redundancy only when they see or terminate plaintext. ReduLink instead assumes the cooperating endpoints already control the plaintext before QUIC protection and can decide whether a reference is safe. This difference matters for HTTP/3, CDN edge delivery, update agents, enterprise VPN clients, and backup clients, where the transport can remain encrypted while the application or endpoint stack performs representation substitution.')
heading(doc, '3.1 Deployment Model', 2)
para(doc, 'The deployment model has three roles: a sender with access to the new byte stream, a receiver with a scoped warm dictionary, and a policy domain that defines whether dictionary state is per connection, per origin, or per tenant. Public Internet mode defaults to per-connection dictionaries. Public artifact mode may use signed or manifest-controlled per-origin dictionaries. Enterprise mode may use per-tenant dictionaries with audit and quota controls. Global private cross-user dictionaries remain out of scope.')
tcap(doc, 'Table 1. ReduLink deployment modes and dictionary scoping.')
add_table(doc, ['Deployment', 'Dictionary scope', 'Why ReduLink rather than rsync/compression'], [
    ['CDN or registry objects', 'Per origin or per client', 'Objects arrive as encrypted streams; file-tree delta negotiation may not be the serving abstraction.'],
    ['Backup/page streams', 'Per endpoint or tenant', 'Page-like chunks recur across snapshots and can be referenced while preserving authenticated reconstruction.'],
    ['Enterprise VPN/admin domain', 'Per tenant', 'Policy can allow shared warm state while keeping middleboxes away from plaintext.'],
    ['Source-tree sync', 'File tree', 'Usually better served by rsync; ReduLink is not targeted here.'],
], widths=[1.4,1.4,3.5], font_size=8)

heading(doc, '4. Protocol and Security Model', 1)
para(doc, 'The sender chunks outgoing bytes and emits FULL frames for new chunks or REF frames for chunks already believed present at the receiver. A REF is deliverable only if dictionary membership, epoch, scope, stream id, reconstructed offset, length, nonce, chunk identity, and authentication tag validate. Otherwise the receiver fails closed or requests semantic FULL repair.')
try:
    doc.add_picture(str(ROOT/'figures/architecture/redulink_architecture.png'), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, 'Figure 1. ReduLink protocol and deployment architecture. The sender chunks outgoing bytes and emits FULL or authenticated REF records carried as compact binary application data inside an encrypted QUIC stream; the receiver validates each reference against dictionary membership and bound context before byte-exact reconstruction, failing closed and requesting semantic FULL repair on any mismatch.')
except Exception:
    pass
para(doc, 'The default privacy scope is per connection. Shared per-origin dictionaries are appropriate only for public artifacts or same-tenant administrative domains. The design excludes global cross-user private dictionaries. The artifact binds frames with an HMAC-based model and an exporter-style key schedule; production QUIC integration should use QUIC TLS exporter material and transport-parameter negotiation.')
tcap(doc, 'Table 2. Threat model: ReduLink controls and corresponding artifact evidence.')
add_table(doc, ['Threat', 'ReduLink control', 'Artifact evidence'], [
    ['Reference forgery', 'HMAC binds epoch, scope, stream, offset, length, nonce, and chunk identity.', 'Tamper and wrong-context tests.'],
    ['Replay', 'Nonce and context binding reject reused frames.', 'Replay rejection tests.'],
    ['Dictionary poisoning', 'Receiver admits only authenticated FULL data or policy-controlled warm state.', 'Dictionary membership and reconstruction tests.'],
    ['Expansion abuse', 'References expand only to negotiated chunk lengths and dictionary entries.', 'Accounting and reconstruction checks.'],
    ['Content-existence leakage', 'No global private cross-user dictionary; scope is explicit policy.', 'Threat model and mode table.'],
], widths=[1.4,2.5,2.3], font_size=8)

heading(doc, '5. Implementation and Metrics', 1)
para(doc, 'The artifact is implemented in Python for reproducibility and uses aioquic for native QUIC stream experiments [18]. ReduLink messages are compact binary application-stream records inside QUIC streams. This exercises real QUIC handshake and encryption, but it does not implement custom QUIC frame types or transport parameters.')
para(doc, 'Three byte layers are reported. Input bytes are the reconstructed application bytes. Stream-payload bytes are the ReduLink or raw application bytes written into QUIC streams. UDP-estimated bytes add a local IPv4/UDP estimate from the loopback proxy path. The paper avoids treating stream-payload multipliers as full wire-rate measurements; they isolate the representation layer so that raw QUIC, ReduLink, gzip, zstd, fixed reuse, and rsync-style baselines can be compared consistently.')
para(doc, 'The compact binary format is included because the JSON baseline was intentionally readable but inflated. The binary format encodes frame type, stream id, reconstructed offset, length, chunk id, nonce, authentication tag, and payload only where required. It is still application-stream data, not QUIC extension-frame data.')
tcap(doc, 'Table 3. Evidence hierarchy: what each layer of evidence supports and does not prove.')
add_table(doc, ['Evidence', 'Supports', 'Does not prove'], [
    ['Model', 'FULL/REF accounting and reconstruction', 'Transport behavior'],
    ['Secure model', 'HMAC binding and replay checks', 'Live QUIC exporter use'],
    ['aioquic stream', 'Encrypted QUIC stream mapping', 'Custom extension frames'],
    ['UDP/IPv4 estimate', 'Local datagram accounting', 'Full packet capture'],
    ['Workloads', 'Workload sensitivity', 'Universal acceleration'],
], widths=[1.3,2.6,2.6])

heading(doc, '6. Evaluation Methodology', 1)
para(doc, 'Evaluation uses four classes of evidence. First, deterministic journal fixtures isolate predicted positive and negative workload shapes. Second, public source-release pairs test ordinary source-tree updates where ReduLink should not be assumed to help. Third, object-aligned public release experiments use the same public bytes but model registry/CDN object delivery rather than tarball synchronization. Fourth, aioquic stream experiments measure the native encrypted stream mapping on positive and negative cases.')
para(doc, 'For each workload, the paper reports ReduLink fixed chunking, ReduLink content-defined chunking where relevant, fixed-block reuse, compression, and rsync where the baseline is structurally applicable. The fixed-block baseline is intentionally strong and simple. If it beats ReduLink, the result is reported rather than hidden, because ReduLink is a security and transport-compatibility layer over reuse, not a superior matching algorithm.')
heading(doc, '7. Workload Results and Baseline Interpretation', 1)
# table warm-state fixtures
def bestcomp(artifact):
    vals=[]
    for method in ['gzip-6','zstd-3']:
        v=mult_for(artifact,method)
        if v: vals.append(v)
    return max(vals) if vals else 0
rows=[]
artifacts=[('scripted-disk-snapshot','disk snap'),('scripted-oci-layer','oci layer'),('scripted-package-metadata','package meta'),('scripted-repository-snapshot','repository snap'),('scripted-structured-logs','logs'),('independent-compressed-negative','compressed neg')]
for art,label in artifacts:
    rl=mult_for(art,'ReduLink fixed') or mult_for(art,'ReduLink CDC') or 0
    # for repository fixed exists and should be shown
    fixed=mult_for(art,'fixed-block-reuse') or 0
    bc=bestcomp(art)
    rows.append([label, fmtx(rl), fmtx(fixed), fmtx(bc), f"{rl-fixed:+.2f}x"])
tcap(doc, 'Table 4. Warm-state workload results (effective stream-payload multiplier): ReduLink versus fixed-block reuse and best compression.')
add_table(doc, ['Workload', 'ReduLink', 'Fixed reuse', 'Best compression', 'RL - fixed'], rows, widths=[1.5,1.1,1.1,1.3,1.1])
para(doc, 'The table is the interpretation guard. On most byte-stable workloads, fixed-block reuse matches or beats ReduLink. The distinct ReduLink contribution is not superior matching; it is authenticated, scoped, stream-compatible reference substitution with explicit repair and privacy rules.')

heading(doc, '7.1 External Public Source Releases', 2)
rows=[]
for r in ext_public:
    label=r['label'].replace('-to-',' to ')
    rs=rsync_public.get(r['label'],{})
    rows.append([label, fmtx(r['redulink_multiplier']), fmtx(r['fixed_block_reuse_multiplier']), fmtx(rs.get('rsync_effective_multiplier_control_plus_data',''))])
tcap(doc, 'Table 5. External public source-release pairs: ReduLink and fixed reuse versus real rsync (negative case for ReduLink).')
add_table(doc, ['Public pair', 'ReduLink', 'Fixed reuse', 'rsync total'], rows, widths=[2.6,1.0,1.0,1.0])
para(doc, 'The independent source-release pairs are negative for ReduLink and strongly positive for rsync. This is an important result. Ordinary related source trees are better served by file-oriented delta transfer than by stream-level reference substitution.')

heading(doc, '7.2 Object-Aligned Public Release Transfers', 2)
rows=[]
for r in ext_obj:
    label=r['label'].split('-object-sequence')[0]
    files=f"{r['unchanged_file_count']}/{r['new_file_count']} unchanged"
    rows.append([label, files, fmtn(r['input_bytes']), fmtx(r['redulink_multiplier']), fmtx(r['secure_multiplier']), fmtx(r['fixed_object_reuse_multiplier']), fmtx(r['gzip_new_object_stream_multiplier'])])
tcap(doc, 'Table 6. Object-aligned public release transfers modeling registry/CDN object delivery.')
add_table(doc, ['Public release', 'File stability', 'Bytes', 'ReduLink', 'Secure RL', 'Fixed object reuse', 'gzip'], rows, widths=[1.2,1.2,0.8,0.8,0.8,1.0,0.8], font_size=7)
para(doc, 'This experiment uses the same public release tarballs but changes the transfer abstraction. Files are treated as object records and chunked independently, modeling registry/CDN object delivery rather than raw tarball or source-tree synchronization. Redis and nginx show positive ReduLink results because many public release objects remain byte-identical; Click is weaker because fewer files are unchanged. This is a stronger external positive signal than a purely synthetic corpus, but it is still not a captured production trace.')
try:
    doc.add_picture(str(ROOT/'figures/external_object_workload/external_object_workload_multipliers.png'), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, 'Figure 2. Object-aligned public release transfers: effective stream-payload multipliers for ReduLink, the authenticated variant, fixed object reuse, and gzip.')
except Exception:
    pass

heading(doc, '7.3 Layer-Like External-Positive Case', 2)
r=ext_pos[0]
tcap(doc, 'Table 7. Layer-like external-positive case derived from public Redis release bytes.')
add_table(doc, ['Case', 'Input bytes', 'ReduLink', 'Secure RL', 'Fixed reuse'], [[r['label'], fmtn(r['new_bytes']), fmtx(r['redulink_multiplier']), fmtx(r['secure_multiplier']), fmtx(r['fixed_block_reuse_multiplier'])]], widths=[2.3,1.0,1.0,1.0,1.0])
para(doc, 'The Redis-layered case is derived from included public Redis release bytes and aligned changed blocks. It is not a production trace; its role is to test the predicted positive case for layer-like objects where stable chunks remain aligned.')

heading(doc, '7.4 Block-Size Sensitivity', 2)
# Build rows for selected artifacts fixed only
arts=['scripted-disk-snapshot','scripted-oci-layer','scripted-repository-snapshot']
rows=[]
for art in arts:
    vals=[]
    for size in ['1024','2048','4096','8192','16384']:
        m=''
        for r in block:
            if r['artifact']==art and r['chunker']=='fixed' and r['chunk_size']==size:
                m=fmtx(r['effective_multiplier']); break
        vals.append(m)
    rows.append([art.replace('scripted-','').replace('-',' ')] + vals)
tcap(doc, 'Table 8. Block-size sensitivity of the effective multiplier (fixed chunking).')
add_table(doc, ['Workload','1 KiB','2 KiB','4 KiB','8 KiB','16 KiB'], rows, widths=[1.5,0.8,0.8,0.8,0.8,0.8])
para(doc, 'The best block size is workload-dependent. The artifact therefore treats 4 KiB as a reproducible default, not as a universal optimum.')
try:
    doc.add_picture(str(ROOT/'figures/block_size/block_size_sensitivity.png'), width=Inches(5.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, 'Figure 4. Block-size sensitivity of the effective multiplier for fixed chunking across three warm-update fixtures.')
except Exception:
    pass

heading(doc, '8. Native QUIC Experiments', 1)
rows=[]
for r in quic_flow:
    rows.append([r['method'].replace('-quic-stream','').replace('redulink-binary','ReduLink').replace('raw','Raw'), r['loss_every'], fmtx(r['effective_multiplier']), fmtx(r['approx_ipv4_udp_multiplier_seen']), 'OK' if r['reconstruction_ok']=='True' else 'FAIL'])
tcap(doc, 'Table 9. Native aioquic stream mapping: raw versus ReduLink under zero and periodic datagram loss.')
add_table(doc, ['Method', 'Loss every', 'Stream x', 'UDP-est x', 'Recon.'], rows, widths=[1.8,0.9,0.9,0.9,0.8])
rows=[]
for r in quic_cases:
    rows.append([r['label'].replace('independent-compressed-negative','compressed negative').replace('external-positive-redis-layered','Redis layered').replace('demo-positive','demo positive'), fmtn(r['input_bytes']), fmtx(r['stream_payload_multiplier']), fmtx(r['approx_ipv4_udp_multiplier_seen']), r['semantic_misses'], 'OK' if r['reconstruction_ok']=='True' else 'FAIL'])
tcap(doc, 'Table 10. Native aioquic stream mapping across positive and negative workload cases.')
add_table(doc, ['Case', 'Input', 'Stream x', 'UDP-est x', 'Misses', 'Recon.'], rows, widths=[1.8,0.9,0.9,0.9,0.7,0.7])
try:
    doc.add_picture(str(ROOT/'figures/quic_workload_cases.png'), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, 'Figure 3. Native aioquic stream mapping on positive and negative workload cases.')
except Exception:
    pass
# Repeated trials summary selected
sel=[]
for r in repeat:
    if r['metric'] in {'raw_client_ms','redulink_stream_multiplier','redulink_udp_est_multiplier','redulink_client_ms'}:
        sel.append([r['metric'].replace('_',' '), r['n'], r['mean'], r.get('stdev',''), r['min'], r['max']])
tcap(doc, 'Table 11. Repeated native QUIC trials (n = 3): run-to-run stability of multipliers and elapsed time.')
add_table(doc, ['Metric', 'n', 'Mean', 'Std. dev.', 'Min', 'Max'], sel, widths=[2.0,0.5,1.0,1.0,1.0,1.0], font_size=7)
para(doc, 'The repeated-trial table is intentionally small but avoids reporting only one favorable QUIC run. The multiplier is deterministic for fixed bytes, while elapsed times vary across runs.')

heading(doc, '9. Component Costs', 1)
components=['fixed_chunking','cdc_chunking','model_fixed_encode_decode','secure_hmac_encode_roundtrip','binary_wire_encode','binary_wire_decode']
rows=[]
for c in components:
    for r in comp:
        if r['component']==c:
            summary=json.loads(r['result_summary'])
            ok=summary.get('reconstruction_ok', summary.get('roundtrip_reconstruction_ok',''))
            if ok is True: ok='OK'
            elif ok is False: ok='FAIL'
            else: ok='not applicable'
            label=c.replace('model_fixed_encode_decode','model fixed roundtrip').replace('secure_hmac_encode_roundtrip','HMAC roundtrip').replace('binary_wire_encode','wire encode').replace('binary_wire_decode','wire decode').replace('fixed_chunking','fixed chunking').replace('cdc_chunking','CDC chunking')
            rows.append([label, fmtn(r['input_bytes']), f"{float(r['throughput_mib_s_local']):.1f}", ok])
tcap(doc, 'Table 12. Component-level cost measurements (local Python artifact timing; not production line-rate).')
add_table(doc, ['Component', 'Input bytes', 'MiB/s local', 'Roundtrip'], rows, widths=[2.1,1.2,1.1,1.2])
para(doc, 'The Python CDC implementation is slow and is not production line-rate evidence. Fixed chunking, HMAC validation, and compact binary encoding are much faster in the local artifact. A production implementation would require optimized native chunking and tighter integration with the transport stack.')

heading(doc, '10. Fairness and Accounting', 1)
para(doc, 'ReduLink must not receive congestion credit for reconstructed bytes. Congestion control and bottleneck accounting count encoded bytes. The package includes wire-byte accounting, QUIC smoke experiments, local UDP/IPv4 byte estimates, and bottleneck emulation. It does not include a full tc/netem or Mininet congestion-control study.')
tcap(doc, 'Table 13. Fairness and accounting summary.')
add_table(doc, ['Metric', 'Value'], [['Fairness rule', 'Congestion counts encoded bytes, not reconstructed bytes'], ['Live congestion-control study', 'Not included'], ['Accounting layer', 'Stream payload and local UDP/IPv4 estimate']])

heading(doc, '11. Why Not Just Compress or Use rsync?', 1)
para(doc, 'Compression is best when repetition is local to one object. rsync is best when both endpoints share a file tree or file object and can run a delta protocol. The ReduLink use case is different: encrypted endpoint streams where dictionary state is already available at the receiver, but the transfer is not naturally a file-tree synchronization session.')
para(doc, 'The evaluation confirms the boundary. On source-release trees, rsync dominates. On package metadata and logs, compression or fixed-block reuse dominates. On object-aligned public releases and layer-like byte-stable objects, ReduLink provides authenticated reference substitution at modest overhead over a simple fixed-reference baseline. The practical motivation is security/context binding and stream compatibility, not beating rsync in rsync-shaped deployments.')
para(doc, 'The motivating cases are therefore not arbitrary file synchronization. They are stream-serving cases in which the endpoint already knows that a receiver has a dictionary state but must still authenticate every reconstruction step and preserve encrypted transport semantics. A CDN edge, registry, backup client, or enterprise service may choose this representation because it avoids exposing plaintext to a WAN optimizer and avoids changing the application protocol into a file-tree delta session.')

heading(doc, '12. Limitations and Future Work', 1)
para(doc, 'The implementation maps ReduLink into QUIC streams rather than adding custom QUIC extension frames or transport parameters. It uses an exporter-style key schedule model rather than live QUIC TLS exporter bytes. The public object-sequence workload is derived from real release bytes, but it is still a transfer-model experiment rather than a captured production trace. The fairness evidence remains conservative and does not replace a full network-emulator congestion study.')
para(doc, 'Future work should add custom QUIC frame integration, live exporter-derived keys, tc/netem or Mininet fairness experiments, packet capture, and independently captured registry or backup traces.')

heading(doc, '13. Reproducibility', 1)
para(doc, 'The package contains source code, public corpora, generated corpora, benchmark scripts, result CSV/JSON files, figures, a citation checker, and tests [15]. The reviewer-facing smoke command is: python3 scripts/run_smoke_validation.py. Full validation is: python3 scripts/run_full_validation.py. Aioquic tests skip gracefully when aioquic is unavailable; installing requirements-dev.txt enables complete QUIC stream validation.')

heading(doc, '14. Conclusion', 1)
para(doc, 'ReduLink is best understood as authenticated, scoped reference substitution for encrypted endpoint streams. Its byte savings are conditional and often reproduced or exceeded by simpler fixed-block reuse, compression, or rsync. The contribution is to make reference substitution explicit, authenticated, privacy-scoped, repairable, and compatible with native QUIC stream transport. The evidence supports ReduLink for warm-state object-aligned or layer-like transfers, while negative source-release results show where it should not be used.')

heading(doc, 'Data and Code Availability', 1)
para(doc, 'All source code, public-corpus manifests, generated corpora, benchmark scripts, result CSV/JSON files, figures, the citation checker, and the test suite are openly available in the ReduLink repository [15] under the MIT License. Every figure and table in this paper is regenerated from the committed result files by the build and validation scripts, so the reported numbers can be reproduced from the artifact. Large external public-release corpora are reconstructed from documented public sources rather than redistributed.')

heading(doc, 'References', 1)
refs=[
"[1] N. T. Spring and D. Wetherall, 'A protocol-independent technique for eliminating redundant network traffic,' ACM SIGCOMM, 2000.",
"[2] A. Anand, V. Sekar, and A. Akella, 'SmartRE: An architecture for coordinated network-wide redundancy elimination,' ACM SIGCOMM, 2009.",
"[3] A. Anand et al., 'Redundancy in network traffic: Findings and implications,' ACM SIGMETRICS, 2009.",
"[4] B. Aggarwal et al., 'EndRE: An end-system redundancy elimination service for enterprises,' USENIX NSDI, 2010.",
"[5] A. Muthitacharoen, B. Chen, and D. Mazieres, 'A low-bandwidth network file system,' ACM SOSP, 2001.",
"[6] J. Iyengar and M. Thomson, 'QUIC: A UDP-based multiplexed and secure transport,' RFC 9000, IETF, 2021.",
"[7] M. Thomson and S. Turner, 'Using TLS to secure QUIC,' RFC 9001, IETF, 2021.",
"[8] J. Iyengar and I. Swett, 'QUIC loss detection and congestion control,' RFC 9002, IETF, 2021.",
"[9] M. Kuehlewind and B. Trammell, 'Applicability of the QUIC transport protocol,' RFC 9308, IETF, 2022.",
"[10] B. Trammell et al., 'Manageability of the QUIC transport protocol,' RFC 9312, IETF, 2022.",
"[11] W. Xia, X. Zou, Y. Zhou, H. Jiang, C. Liu, D. Feng, Y. Hua, Y. Hu, and Y. Zhang, 'The design of fast content-defined chunking for data deduplication based storage systems,' IEEE Transactions on Parallel and Distributed Systems, vol. 31, no. 9, 2020.",
"[12] M. Gregoriadis, L. Balduf, B. Scheuermann, and J. Pouwelse, 'A thorough investigation of content-defined chunking algorithms for data deduplication,' arXiv:2409.06066, 2024.",
"[13] B. Alexeev, C. Percival, and Y. X. Zhang, 'Chunking attacks on file backup services using content-defined chunking,' arXiv:2504.02095, 2025.",
"[14] IEEE 802.3 Ethernet Working Group, 'IEEE 802.3 Ethernet Working Group active projects,' accessed June 2026.",
"[15] M. Nguyen, 'ReduLink artifact and reproducibility package,' GitHub repository, 2026. https://github.com/pinkysworld/redulink-deduplex-quic",
"[16] D. Harnik, B. Pinkas, and A. Shulman-Peleg, 'Side channels in cloud services: Deduplication in cloud storage,' IEEE Security and Privacy, 2010.",
"[17] M. Bellare, S. Keelveedhi, and T. Ristenpart, 'DupLESS: Server-aided encryption for deduplicated storage,' USENIX Security, 2013.",
"[18] aioquic project contributors, 'aioquic: QUIC and HTTP/3 implementation in Python,' software repository, version 1.3.0, 2026. https://github.com/aiortc/aioquic"
]
for ref in refs:
    para(doc, ref, size=8, after=1)

# save
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.author = 'Michél Nguyen'
doc.core_properties.last_modified_by = 'Michél Nguyen'
doc.core_properties.subject = 'ORCID: 0000-0001-6834-4422; University of the People'
doc.core_properties.comments = 'ReduLink journal-ready v2.8 manuscript; ORCID: 0000-0001-6834-4422; University of the People'
doc.save(OUT)
print(OUT)
