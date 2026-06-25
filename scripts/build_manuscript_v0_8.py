#!/usr/bin/env python3
"""Build the v0.8 peer-review-strengthened manuscript DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "paper/submission/ReduLink_Deduplex_QUIC_full_draft_v0_7_public_corpora_rsync_prototype.docx"
OUT = ROOT / "paper/submission/ReduLink_Deduplex_QUIC_full_draft_v0_8_peer_review_strengthened.docx"
FIGURE = ROOT / "figures/effective_multiplier_by_workload.png"


def norm(text: str) -> str:
    return " ".join(text.split())


def find_paragraph(doc: Document, startswith: str):
    for paragraph in doc.paragraphs:
        if norm(paragraph.text).startswith(startswith):
            return paragraph
    raise ValueError(f"paragraph not found: {startswith}")


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    new = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new._p)
    if style:
        new.style = style
    set_text(new, text)
    return new


def remove_block(block) -> None:
    element = block._element
    element.getparent().remove(element)


def insert_table_after(paragraph, rows: list[list[str]], style: str = "Table Grid"):
    doc = paragraph.part.document
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = style
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
    paragraph._p.addnext(table._tbl)
    return table


def compact_table(table, font_size: int = 8) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def replace_table_after_caption(doc: Document, caption_prefix: str, rows: list[list[str]]) -> None:
    caption = find_paragraph(doc, caption_prefix)
    for table in doc.tables:
        prev = table._tbl.getprevious()
        if prev is not None and prev == caption._p:
            remove_block(table)
            break
    insert_table_after(caption, rows)


def delete_paragraphs_starting(doc: Document, prefixes: list[str]) -> None:
    for paragraph in list(doc.paragraphs):
        text = norm(paragraph.text)
        if any(text.startswith(prefix) for prefix in prefixes):
            remove_block(paragraph)


def main() -> None:
    doc = Document(SRC)

    set_text(
        doc.paragraphs[1],
        "Michél Nguyen Independent Researcher | minh.systems Draft version 0.8 peer-review-strengthened evidence, transport-semantics, and security revision, June 2026",
    )
    set_text(
        find_paragraph(doc, "Abstract."),
        "Abstract. Ethernet and wide-area transport protocols are typically evaluated by their physical or nominal line rate. However, many real traffic classes contain repeated objects, pages, layers, templates, or versioned artifacts. ReduLink is a security-scoped protocol design for authenticated reference substitution in cooperative encrypted endpoints. Rather than increasing physical line rate, ReduLink increases effective reconstructed payload throughput when a receiver can validate compact references against an epoch-scoped dictionary. This draft specifies the ReduLink abstraction, a Deduplex-QUIC extension-frame profile, sender and receiver state machines, security properties, privacy modes, and conservative accounting rules. The accompanying artifact validates byte-exact reconstruction, fail-closed reference handling, public-corpora benchmarks, a fixed-block reuse approximation baseline, local cost columns, and a TCP endpoint-reconstruction prototype. Production cryptographic binding, replay windows, real QUIC integration, congestion-fairness experiments, and cross-tenant privacy enforcement remain implementation requirements.",
    )

    for paragraph in doc.paragraphs:
        text = norm(paragraph.text)
        if text.startswith("A reproducible offline simulator and repository package"):
            set_text(
                paragraph,
                "A reproducible artifact package with fixed and content-defined chunking, public-corpora fetch scripts, fixed-block reuse baseline, local CPU/RSS cost columns, socket prototype, automated tests, CI, CSV outputs, and generated figures.",
            )
        elif text == "2.6 Gap in the literature":
            gap_heading = paragraph
        elif text.startswith("The reproducibility package accompanying this draft now includes"):
            set_text(
                paragraph,
                "The reproducibility package accompanying this draft includes the Python model, unit tests, GitHub Actions workflow, selected earlier artifact CSV, fetched public-corpora fixture, fixed-block reuse approximation, benchmark scripts with local cost columns, plot-generation script, generated figures, minimal socket prototype, threat model, and RFC-style protocol appendix. Larger validation should still run on frozen OCI layers, Linux kernel tarballs, git pack snapshots, package metadata, and VM/container snapshots before any production-scale claim.",
            )
        elif text.startswith("The paper should no longer frame"):
            set_text(
                paragraph,
                "The evaluation is separated into evidence levels: analytic model, simulator, controlled synthetic workloads, frozen public-corpora fixture, and endpoint reconstruction prototype. This separation prevents synthetic multipliers from being read as production trace results and makes the remaining validation gap explicit.",
            )
        elif text.startswith("The baseline table compares raw bytes"):
            set_text(
                paragraph,
                "The baseline table compares raw bytes, gzip, zstd when available, a fixed-block reuse approximation inspired by rsync-family delta transfer, ReduLink with fixed chunking, ReduLink with content-defined chunking, and composition cases where compression is applied before or after ReduLink. The appropriate claim is that ReduLink is not a replacement for compression or rsync: compression exploits redundancy inside the current object, delta transfer is strong for file updates, and ReduLink suppresses repeated chunks across endpoint-controlled dictionary history.",
            )
        elif text.startswith("Benchmark reproduction is intentionally command-based"):
            set_text(
                paragraph,
                "Benchmark reproduction is command-based: python3 benchmarks/fetch_public_corpora.py fetches pinned public pairs; bash benchmarks/run_synthetic_suite.sh emits results/synthetic_suite.csv; bash benchmarks/run_public_artifacts.sh --manifest benchmarks/public_artifacts_manifest.csv emits results/public_artifact_suite.csv; and python3 scripts/plot_results.py results/synthetic_suite.csv --output-dir figures regenerates figures. The CSV metadata records Python, platform, gzip/zstd availability, and the meaning of local cost columns.",
            )
        elif text.startswith("The repository now includes tests for byte-exact FULL/REF reconstruction"):
            set_text(
                paragraph,
                "The repository now includes tests for byte-exact FULL/REF reconstruction, random-data negative controls, warm-dictionary gains, safe REF-miss failure, public-manifest validation, and both fixed and content-defined chunkers. GitHub Actions runs the test suite and a benchmark smoke test on push and pull request. This gives readers a concrete correctness story rather than only a prose protocol description.",
            )
        elif text.startswith("The repository now also includes a minimal localhost socket prototype"):
            set_text(
                paragraph,
                "The repository now also includes a minimal localhost socket prototype. In demo mode, a client sends modeled FULL/REF frames over TCP to a receiver, which reconstructs the update byte-exactly from transmitted frames and warm dictionary. The prototype validates endpoint cooperation and reconstruction only; it does not exercise QUIC packet protection, ACK/loss recovery, stream final-size handling, flow control, congestion fairness, 0-RTT, migration, or extension-frame parsing.",
            )
        elif text.startswith("The v0.5 benchmark suite therefore reports"):
            set_text(
                paragraph,
                "The benchmark suite therefore reports standalone compression baselines, a fixed-block reuse approximation, ReduLink fixed/CDC rows, and compression composition cases. This makes the distinction measurable: if zstd or gzip already removes repeated byte identity, ReduLink should not claim additional gain; if update-like artifacts preserve repeated chunk identity across history, ReduLink can reduce transmitted bytes even though congestion control still accounts only for wire bytes.",
            )
        elif text.startswith("The main reviewer objection will be"):
            set_text(
                paragraph,
                "The scope of claim is deliberately narrow: ReduLink is not a faster physical link, not transparent middlebox optimization, not universal compression, and not yet a QUIC implementation. Its defended contribution is an endpoint-controlled, authenticated reference-substitution design for redundant warm-state workloads, with explicit privacy scope, miss repair, expansion bounds, flow-control semantics, and reproducible artifact evidence.",
            )
        elif text.startswith("The current evaluation combines analytic modeling"):
            set_text(
                paragraph,
                "The current evaluation separates analytic modeling, simulation, small frozen public-corpora fixtures, and a TCP endpoint prototype; it still requires larger public corpora and real QUIC implementation experiments before production-scale claims.",
            )
        elif text.startswith("To strengthen the manuscript beyond an analytic bandwidth model"):
            set_text(
                paragraph,
                "To strengthen the manuscript beyond an analytic bandwidth model, this revision treats the repository as part of the evidence. The current artifact remains an endpoint-controlled redundancy-suppressed representation layer with epoch dictionaries, chunk identifiers, reference frames, bounded dictionary capacity, reference-miss fallback, and explicit wire-byte accounting. The repository now adds tests, baseline comparison commands, public-artifact hooks, local cost columns, and CSV-driven figure generation. The simulator is still not presented as a population estimate for the entire Internet; it is a controlled proof-of-concept showing when the protocol mechanism can and cannot create effective throughput gain.",
            )
        elif text.startswith("The v0.7 reproducibility package contains"):
            set_text(
                paragraph,
                "The v0.8 reproducibility package contains the runnable model, tests, GitHub Actions workflow, selected earlier artifact CSV, synthetic CSV, public-corpora manifest and result CSV, fixed-block reuse approximation, benchmark scripts with local cost columns, plot-generation script, regenerated figures, minimal socket prototype, threat model, and RFC-style protocol appendix. Core files include src/redulink_proto_v0_5.py, tests/, benchmarks/, prototypes/, scripts/plot_results.py, results/, figures/, docs/protocol_summary.md, and docs/threat_model.md. Representative commands are:",
            )
        elif text.startswith("bash benchmarks/run_public_artifacts.sh ubuntu-base="):
            set_text(
                paragraph,
                "bash benchmarks/run_public_artifacts.sh --manifest benchmarks/public_artifacts_manifest.csv",
            )
        elif text.startswith("ReduLink reframes network acceleration"):
            set_text(
                paragraph,
                "ReduLink reframes network acceleration from physical-speed escalation to effective reconstructed payload throughput. A 1 Gbit/s link cannot physically transmit more than its line rate, but a receiver can reconstruct more than 1 Gbit/s of original payload when repeated bytes are replaced by authenticated references. The proposed protocol makes this idea compatible with modern encrypted WANs by using endpoint negotiation, epoch-scoped dictionaries, authenticated references, bounded expansion, reference-miss fallback, and congestion accounting based on wire bytes. The expected gains are workload-dependent and carry CPU, memory, and privacy costs. The resulting research contribution is not a faster Ethernet PHY, but a security-scoped transport-layer or link-adjacent primitive for effective bandwidth expansion under redundancy.",
            )

    comparison_rows = [
        ["System", "Placement", "E2E encrypted traffic", "Auth refs / epochs", "QUIC semantics", "Privacy posture"],
        ["Spring/Wetherall RE", "Middlebox/network", "Weak for QUIC/TLS", "No / no", "No", "Network-visible payloads"],
        ["LBFS / rsync", "File/application", "Yes, file-specific", "Partial / no", "No", "Application-specific"],
        ["EndRE", "Endpoint service", "Better than middleboxes", "Limited / unclear", "No", "Enterprise endpoint scope"],
        ["SmartRE", "Coordinated middleboxes", "Weak for E2E QUIC", "No / no", "No", "Enterprise/provider scope"],
        ["ReduLink", "Endpoint/QUIC/VPN", "Yes, endpoint-negotiated", "Yes / yes", "Specified", "Per-connection default; no global cross-user mode"],
    ]
    gap_para = find_paragraph(doc, "2.6 Gap in the literature")
    insert_after = gap_para
    for candidate in doc.paragraphs:
        if candidate._p.getprevious() is gap_para._p:
            insert_after = candidate
            break
    caption = insert_paragraph_after(insert_after, "Table 1a. Related-work comparison and ReduLink gap.", "Caption")
    compact_table(insert_table_after(caption, comparison_rows), 8)

    sec_para = find_paragraph(doc, "7.1 Threat model")
    sec_body = sec_para
    for candidate in doc.paragraphs:
        if candidate._p.getprevious() is sec_para._p:
            sec_body = candidate
            break
    security_caption = insert_paragraph_after(sec_body, "Table 5a. Security properties, mechanisms, and artifact status.", "Caption")
    compact_table(insert_table_after(
        security_caption,
        [
            ["Property", "Claim", "Required mechanism", "Current artifact status"],
            ["Integrity", "Output equals input or fails closed.", "FULL/REF authentication, chunk-id, offset, and length binding.", "Reconstruction and mismatch tests; crypto not implemented."],
            ["Context binding", "REF cannot replay across connection, epoch, stream, offset, origin, or scope.", "Exporter-derived keys, epoch id, stream id, offset, dictionary id, nonce/replay window.", "Specified, not implemented."],
            ["Dictionary safety", "Only authenticated FULL or signed-manifest chunks are admitted.", "Authenticated FULL, manifest commitment, admission and eviction policy.", "Chunk-id checks modeled; manifest policy pending."],
            ["Expansion bound", "Small REF cannot trigger unbounded work or delivery.", "Per-frame, per-stream, and per-epoch reconstructed-byte caps.", "Basic accounting modeled; QUIC flow-control enforcement pending."],
            ["Privacy scope", "No private cross-user possession oracle in public mode.", "Per-connection default; no global cross-user dictionary.", "Policy specified; enforcement pending."],
        ],
    ), 8)

    side_para = find_paragraph(doc, "Deduplication can leak information")
    side_after = insert_paragraph_after(
        side_para,
        "Like other deduplication systems, ReduLink can create a content-existence oracle when an adversary can choose payloads or references and observe wire size, timing, MISS count, fallback FULL size, DICT_ACK behavior, or application latency. Per-connection dictionaries remove the cross-user oracle in public-WAN mode, but they do not remove all leakage inside one connection, one origin, one tenant, or one administrative domain. Shared dictionaries are therefore limited to public artifacts, same-tenant deployments, or explicitly accepted policy domains.",
    )
    privacy_caption = insert_paragraph_after(side_after, "Table 5b. Privacy modes and required controls.", "Caption")
    compact_table(insert_table_after(
        privacy_caption,
        [
            ["Mode", "Allowed dictionary scope", "Leakage risk", "Default?", "Required controls"],
            ["Public Internet", "Per-connection only", "Same-connection access-pattern leakage", "Yes", "No cross-user refs, 1-RTT only, bounded epochs"],
            ["Public artifacts", "Per-origin signed manifest", "Artifact-version inference", "Optional", "Public-only content, manifest commitment, expiration"],
            ["Enterprise VPN", "Tenant/admin domain", "Intra-tenant content-existence leakage", "Optional", "Tenant policy, quotas, audit, opt-out"],
            ["Global cross-user", "Any user", "Private content-possession leakage", "No", "Out of scope"],
        ],
    ), 8)

    set_text(find_paragraph(doc, "Table 4. Simulator parameters"), "Table 4. Simulator parameters used in the proof-of-concept evaluation.")
    set_text(find_paragraph(doc, "Table 6. v0.7"), "Table 6. v0.8 evidence excerpt from generated CSV outputs: evidence level, fixed-block reuse approximation, corrected public-corpora rows, and local cost columns.")
    set_text(find_paragraph(doc, "Figure 6."), "Figure 6. Regenerated v0.8 synthetic-suite figure: effective reconstructed throughput multiplier by workload and method, including gzip, zstd, fixed-block reuse, ReduLink fixed, and ReduLink CDC.")
    set_text(find_paragraph(doc, "11.4 Reviewer risks"), "11.4 Scope of claims")

    table6_rows = [
        ["Dataset", "Evidence level", "Method", "Multiplier", "Saving", "CPU ms"],
        ["logs", "synthetic warm update", "fixed-block reuse", "72.694x", "0.986", "11.113"],
        ["logs", "synthetic warm update", "ReduLink fixed", "63.631x", "0.984", "7.666"],
        ["mixed", "synthetic warm update", "ReduLink fixed", "4.890x", "0.796", "8.735"],
        ["nginx-changes", "public changed pair", "fixed-block reuse", "72.956x", "0.986", "3.295"],
        ["nginx-changes", "public changed pair", "ReduLink CDC", "53.126x", "0.981", "621.301"],
        ["redis-readme", "public changed pair", "ReduLink CDC", "1.565x", "0.361", "18.408"],
        ["cpython-http-server", "public changed pair", "ReduLink CDC", "0.998x", "0.000", "48.086"],
        ["linux-kernel-parameters", "public changed pair", "ReduLink CDC", "0.998x", "0.000", "258.930"],
        ["random control", "negative control", "ReduLink fixed", "0.997x", "0.000", "n/a"],
    ]
    replace_table_after_caption(doc, "Table 6.", table6_rows)
    compact_table(doc.tables[-1], 8)

    # Replace the old Figure 6 image paragraph with the regenerated plot.
    for paragraph in list(doc.paragraphs):
        text = norm(paragraph.text)
        if paragraph._p.xpath(".//w:drawing") and paragraph._p.getprevious() is find_paragraph(doc, "Figure 6.")._p:
            remove_block(paragraph)
            break
    fig_para = insert_paragraph_after(find_paragraph(doc, "Figure 6."), "")
    fig_para.add_run().add_picture(str(FIGURE), width=Inches(6.4))

    for paragraph in doc.paragraphs:
        if "ReduLink / Deduplex-QUIC v0.5 repository" in paragraph.text:
            set_text(
                paragraph,
                '[14] M. Nguyen, "ReduLink / Deduplex-QUIC repository and reproducibility package," GitHub repository, 2026. Available: https://github.com/pinkysworld/redulink-deduplex-quic',
            )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "v0.5 simulator REF metadata estimate" in cell.text:
                    cell.text = cell.text.replace(
                        "v0.5 simulator REF metadata estimate",
                        "Current artifact REF metadata estimate",
                    )
    refs = [norm(p.text) for p in doc.paragraphs]
    ref14 = find_paragraph(doc, "[14] M. Nguyen")
    if not any("Harnik" in r and "Side channels" in r for r in refs):
        ref15 = insert_paragraph_after(
            ref14,
            '[15] D. Harnik, B. Pinkas, and A. Shulman-Peleg, "Side channels in cloud services: Deduplication in cloud storage," IEEE Security & Privacy, 2010.',
        )
        insert_paragraph_after(
            ref15,
            '[16] M. Bellare, S. Keelveedhi, and T. Ristenpart, "DupLESS: Server-aided encryption for deduplicated storage," USENIX Security, 2013.',
        )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
